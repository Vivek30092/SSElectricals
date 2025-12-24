from django.shortcuts import render, redirect, get_object_or_404
from django.core.files.storage import FileSystemStorage
from django.contrib.admin.views.decorators import staff_member_required # Keep for reference or fallback
from .decorators import admin_required, staff_required
from django.contrib import messages
from .models import Appointment, AdminActivityLog, AdminSession, Order, Product, Category, ProductImage, OrderItem, Review, DailySales, DailyExpenditure, PurchaseEntry, CustomUser, Notification, UserNotification, ServicePrice, Electrician, Warranty, ServiceType
import os
from django.conf import settings
import pandas as pd
from django.db.models import Sum, Count, F, Q
from django.db import models
from django.db.models.functions import TruncDate, TruncMonth, Coalesce
from django.utils import timezone
from .forms import ProductForm, ReviewForm, DailySalesForm, DailyExpenditureForm
from django.core.mail import send_mail
from django.http import HttpResponse, JsonResponse
from django.core.paginator import Paginator
import csv
import datetime
import json
from datetime import timedelta
from dateutil.relativedelta import relativedelta
from django.utils import timezone
import numpy as np

def get_date_range_filter(range_type, custom_start=None, custom_end=None):
    """Calculate date range based on filter type."""
    today = timezone.now().date()
    
    range_map = {
        'this_month': (today.replace(day=1), today),
        'last_month': (
            (today.replace(day=1) - timedelta(days=1)).replace(day=1),
            today.replace(day=1) - timedelta(days=1)
        ),
        'last_3_months': (today - relativedelta(months=3), today),
        'last_6_months': (today - relativedelta(months=6), today),
        'last_year': (today - relativedelta(years=1), today),
        'last_2_years': (today - relativedelta(years=2), today),
        'last_3_years': (today - relativedelta(years=3), today),
    }
    
    if range_type == 'custom' and custom_start and custom_end:
        return (custom_start, custom_end)
    
    return range_map.get(range_type, (None, None))


@staff_required # Updated to custom decorator
def admin_dashboard(request):
    """
    Admin Dashboard - Financial Metrics
    
    CRITICAL BUSINESS RULE (Non-Negotiable):
    ========================================
    All financial metrics (revenue, sales, profit) are calculated EXCLUSIVELY from:
      - Manual DailySales entries
      - Manual DailyExpenditure entries
      - Manual PurchaseEntry records
    
    Orders and deliveries are OPERATIONAL records only.
    They do NOT affect financial calculations under any circumstances.
    
    This ensures:
      - Clean separation between operations and accounting
      - Accurate real-world financial tracking
      - No accidental double-counting
      - Admin retains full control of business numbers
    """
    
    # ==================================================================================
    # FINANCIAL METRICS - MANUAL ENTRY ONLY (NO ORDER DATA)
    # ==================================================================================
    
    # Total Revenue: ONLY from manually entered daily sales
    total_revenue = DailySales.objects.aggregate(Sum('total_sales'))['total_sales__sum'] or 0
    
    # Cash vs Online breakdown from manual sales entries
    total_cash_received = DailySales.objects.aggregate(Sum('cash_received'))['cash_received__sum'] or 0
    total_online_received = DailySales.objects.aggregate(Sum('online_received'))['online_received__sum'] or 0


    # Expenses: From manual daily expenditure entries (using combined online + cash model)
    total_expenses = DailyExpenditure.objects.aggregate(Sum('total'))['total__sum'] or 0
    online_expenses = DailyExpenditure.objects.aggregate(Sum('online_amount'))['online_amount__sum'] or 0
    cash_expenses = DailyExpenditure.objects.aggregate(Sum('cash_amount'))['cash_amount__sum'] or 0
    
    # Purchases (Inventory Cost): From manual purchase entries
    total_purchases = PurchaseEntry.objects.aggregate(Sum('total_cost'))['total_cost__sum'] or 0

    # Net Profit Calculation: Revenue - Expenses - Purchases (NO ORDER DATA)
    profit_loss = float(total_revenue) - float(total_expenses) - float(total_purchases)

    # ==================================================================================
    # OPERATIONAL METRICS (FOR REFERENCE ONLY - NOT USED IN FINANCIAL CALCULATIONS)
    # ==================================================================================
    
    total_orders = Order.objects.count()
    pending_orders = Order.objects.filter(status='Pending').count()
    pending_appointments = Appointment.objects.filter(status='Pending').count()

    # ==================================================================================
    # SALES TREND CHART DATA (FROM MANUAL DAILY SALES ONLY)
    # ==================================================================================
    
    today = timezone.now().date()
    
    # Date Range Filter logic
    range_option = request.GET.get('range', '30') # Default to 30 days
    start_date = None
    
    if range_option == '30':
        start_date = today - datetime.timedelta(days=30)
    elif range_option == '90':
         start_date = today - datetime.timedelta(days=90)
    elif range_option == '180':
         start_date = today - datetime.timedelta(days=180)
    elif range_option == '270':
         start_date = today - datetime.timedelta(days=270)
    elif range_option == '365':
         start_date = today - datetime.timedelta(days=365)
    elif range_option == '730':
         start_date = today - datetime.timedelta(days=730)
    elif range_option == '1095':
         start_date = today - datetime.timedelta(days=1095)
    elif range_option == 'all':
         start_date = None # All time
    else:
        start_date = today - datetime.timedelta(days=30)
    
    # Fetch DailySales for chart
    sales_qs = DailySales.objects.all()
    if start_date:
        sales_qs = sales_qs.filter(date__gte=start_date)
        
    daily_sales_data = sales_qs.values('date') \
        .annotate(sales=Sum('total_sales')) \
        .order_by('date')

    dates = [x['date'].strftime('%Y-%m-%d') for x in daily_sales_data]
    sales = [float(x['sales'] or 0) for x in daily_sales_data]

    # Fetch DailyExpenditure for comparison chart
    expenses_qs = DailyExpenditure.objects.all()
    if start_date:
        expenses_qs = expenses_qs.filter(date__gte=start_date)
    
    daily_expenses_data = expenses_qs.values('date') \
        .annotate(expense=Sum('total')) \
        .order_by('date')
    
    # Create a dictionary for quick lookup
    expenses_dict = {x['date'].strftime('%Y-%m-%d'): float(x['expense'] or 0) for x in daily_expenses_data}
    
    # Match expenses to sales dates (fill with 0 if no expense for that date)
    expenses = [expenses_dict.get(date, 0) for date in dates]

    # Category Statistics (for product management reference)
    category_stats = Product.objects.values('category__name').annotate(count=Count('id')).order_by('-count')
    cat_labels = [x['category__name'] or 'Uncategorized' for x in category_stats]
    cat_data = [x['count'] for x in category_stats]

    # Recent operational records (for quick reference)
    recent_orders = Order.objects.select_related('user').order_by('-created_at')[:10]
    recent_appointments = Appointment.objects.order_by('-created_at')[:10]

    context = {
        # Financial Metrics (Manual Entry Only)
        'total_revenue': total_revenue,
        'total_cash_received': total_cash_received,
        'total_online_received': total_online_received,
        'total_expenses': total_expenses,
        'online_expenses': online_expenses,
        'cash_expenses': cash_expenses,
        'total_purchases': total_purchases,
        'profit_loss': profit_loss,
        
        # Operational Metrics (Not used in financial calculations)
        'total_orders': total_orders,
        'pending_orders': pending_orders,
        'pending_appointments': pending_appointments,
        
        # Chart Data (From Manual Sales Only)
        'dates_json': json.dumps(dates),
        'sales_json': json.dumps(sales),
        'expenses_json': json.dumps(expenses),
        'cat_labels_json': json.dumps(cat_labels),
        'cat_data_json': json.dumps(cat_data),
        
        # Quick Reference
        'recent_orders': recent_orders,
        'recent_appointments': recent_appointments,
        'current_range': range_option,
    }
    return render(request, 'admin/admin_dashboard.html', context)

@staff_required
def admin_analytics_new(request):
    """Enhanced Analytics Dashboard with Dynamic Filters"""
    from django.db.models import Sum
    from django.db.models.functions import TruncMonth, TruncQuarter, TruncYear
    import json
    from firstApp.models import DailySales, DailyExpenditure
    
    # Get filter parameters
    quarterly_year = request.GET.get('quarterly_year', str(timezone.now().year))
    monthly_range = request.GET.get('monthly_range', 'last_6_months')
    comparison_range = request.GET.get('comparison_range', 'last_6_months')
    expense_range = request.GET.get('expense_range', 'last_6_months')
    
    # Helper function
    def safe_float(val):
        return float(val) if val else 0.0
    
    # Get all available years from sales data
    available_years = DailySales.objects.dates('date', 'year', order='DESC')
    available_years_list = [d.year for d in available_years]
    
    # ===== OVERALL TOTALS (No filtering) =====
    sales_agg = DailySales.objects.aggregate(
        total=Sum('total_sales'),
        online=Sum('online_received'),
        cash=Sum('cash_received'),
        labor=Sum('labor_charge'),
        delivery=Sum('delivery_charge')
    )
    
    expense_agg = DailyExpenditure.objects.aggregate(
        total=Sum('total'),
        online=Sum('online_amount'),
        cash=Sum('cash_amount')
    )
    
    sales_count = DailySales.objects.count()
    expense_count = DailyExpenditure.objects.count()
    
    total_sales = safe_float(sales_agg['total'])
    total_online_sales = safe_float(sales_agg['online'])
    total_cash_sales = safe_float(sales_agg['cash'])
    total_labor = safe_float(sales_agg['labor'])
    total_delivery = safe_float(sales_agg['delivery'])
    
    total_expense = safe_float(expense_agg['total'])
    total_online_exp = safe_float(expense_agg['online'])
    total_cash_exp = safe_float(expense_agg['cash'])
    
    avg_sales = (total_sales / sales_count) if sales_count > 0 else 0.0
    avg_expense = (total_expense / expense_count) if expense_count > 0 else 0.0
    
    online_pct = (total_online_sales / total_sales * 100) if total_sales > 0 else 0
    cash_pct = (total_cash_sales / total_sales * 100) if total_sales > 0 else 0
    online_exp_pct = (total_online_exp / total_expense * 100) if total_expense > 0 else 0
    cash_exp_pct = (total_cash_exp / total_expense * 100) if total_expense > 0 else 0
    
    # ===== QUARTERLY PERFORMANCE (Filtered by Year) =====
    try:
        selected_year = int(quarterly_year)
    except:
        selected_year = timezone.now().year
    
    quarterly_sales = DailySales.objects.filter(
        date__year=selected_year
    ).annotate(
        quarter=TruncQuarter('date')
    ).values('quarter').annotate(
        total=Sum('total_sales')
    ).order_by('quarter')
    
    quarterly_list = list(quarterly_sales)
    quarterly_labels = []
    quarterly_values = []
    for q in quarterly_list:
        quarter_date = q['quarter']
        quarter_num = (quarter_date.month - 1) // 3 + 1
        label = f"Q{quarter_num}"
        quarterly_labels.append(label)
        quarterly_values.append(safe_float(q['total']))
    
    # ===== MONTHLY SALES & GROWTH (Filtered by Range) =====
    monthly_start, monthly_end = get_date_range_filter(monthly_range)
    monthly_qs = DailySales.objects.all()
    if monthly_start and monthly_end:
        monthly_qs = monthly_qs.filter(date__range=[monthly_start, monthly_end])
    
    monthly_sales = monthly_qs.annotate(
        month=TruncMonth('date')
    ).values('month').annotate(
        total=Sum('total_sales')
    ).order_by('month')
    
    monthly_sales_list = list(monthly_sales)
    monthly_sales_labels = [m['month'].strftime('%b %Y') for m in monthly_sales_list]
    monthly_sales_values = [safe_float(m['total']) for m in monthly_sales_list]
    
    # Calculate growth
    if len(monthly_sales_values) >= 2:
        latest_month = monthly_sales_values[-1]
        previous_month = monthly_sales_values[-2]
        monthly_growth = ((latest_month - previous_month) / previous_month * 100) if previous_month > 0 else 0
    else:
        monthly_growth = 0
    
    # ===== SALES VS EXPENSES COMPARISON (Filtered by Range) =====
    comp_start, comp_end = get_date_range_filter(comparison_range)
    
    sales_comp_qs = DailySales.objects.all()
    expenses_comp_qs = DailyExpenditure.objects.all()
    
    if comp_start and comp_end:
        sales_comp_qs = sales_comp_qs.filter(date__range=[comp_start, comp_end])
        expenses_comp_qs = expenses_comp_qs.filter(date__range=[comp_start, comp_end])
    
    sales_comparison = sales_comp_qs.annotate(
        month=TruncMonth('date')
    ).values('month').annotate(
        total=Sum('total_sales')
    ).order_by('month')
    
    expenses_comparison = expenses_comp_qs.annotate(
        month=TruncMonth('date')
    ).values('month').annotate(
        total=Sum('total')
    ).order_by('month')
    
    comparison_labels = [s['month'].strftime('%b %Y') for s in sales_comparison]
    comparison_sales = [safe_float(s['total']) for s in sales_comparison]
    
    expenses_dict = {e['month']: safe_float(e['total']) for e in expenses_comparison}
    comparison_expenses = [expenses_dict.get(s['month'], 0) for s in sales_comparison]
    
    # ===== DAILY EXPENSE TREND (Filtered by Range) =====
    exp_start, exp_end = get_date_range_filter(expense_range)
    expense_trend_qs = DailyExpenditure.objects.all()
    
    if exp_start and exp_end:
        expense_trend_qs = expense_trend_qs.filter(date__range=[exp_start, exp_end])
    
    daily_expenses = expense_trend_qs.order_by('date').values('date', 'total')
    expense_dates = [e['date'].strftime('%Y-%m-%d') for e in daily_expenses]
    expense_values = [safe_float(e['total']) for e in daily_expenses]
    
    # ===== INSIGHTS (Min/Max from overall data) =====
    min_sale_day = DailySales.objects.order_by('total_sales').first()
    max_sale_day = DailySales.objects.order_by('-total_sales').first()
    
    # Best month from all time
    all_monthly_sales = DailySales.objects.annotate(
        month=TruncMonth('date')
    ).values('month').annotate(
        total=Sum('total_sales')
    ).order_by('-total')
    
    if all_monthly_sales:
        best_month = all_monthly_sales[0]
        max_sales_month = best_month['month'].strftime('%B %Y')
        max_sales_month_value = safe_float(best_month['total'])
    else:
        max_sales_month = 'N/A'
        max_sales_month_value = 0
    
    # ===== WEEKDAY PERFORMANCE =====
    weekday_sales = DailySales.objects.values('day').annotate(
        total=Sum('total_sales')
    ).order_by('day')
    
    weekday_labels = [w['day'] for w in weekday_sales]
    weekday_values = [safe_float(w['total']) for w in weekday_sales]
    
    # ===== YEARLY TREND =====
    yearly_sales = DailySales.objects.annotate(
        year=TruncYear('date')
    ).values('year').annotate(
        total=Sum('total_sales')
    ).order_by('year')
    
    yearly_labels = [y['year'].strftime('%Y') for y in yearly_sales]
    yearly_values = [safe_float(y['total']) for y in yearly_sales]
    
    context = {
        # Overall stats
        'total_sales': total_sales,
        'total_expense': total_expense,
        'total_online_sales': total_online_sales,
        'total_cash_sales': total_cash_sales,
        'total_labor': total_labor,
        'total_delivery': total_delivery,
        'total_online_exp': total_online_exp,
        'total_cash_exp': total_cash_exp,
        'avg_sales': avg_sales,
        'avg_expense': avg_expense,
        'online_sales_percentage': round(online_pct, 1),
        'cash_sales_percentage': round(cash_pct, 1),
        'online_expense_percentage': round(online_exp_pct, 1),
        'cash_expense_percentage': round(cash_exp_pct, 1),
        
        # Insights
        'min_sale_day': min_sale_day,
        'max_sale_day': max_sale_day,
        'max_sales_month': max_sales_month,
        'max_sales_month_value': max_sales_month_value,
        
        # Filter states
        'selected_quarterly_year': selected_year,
        'available_years': available_years_list,
        'selected_monthly_range': monthly_range,
        'selected_comparison_range': comparison_range,
        'selected_expense_range': expense_range,
        
        # Chart data
        'quarterly_labels_json': json.dumps(quarterly_labels),
        'quarterly_values_json': json.dumps(quarterly_values),
        'monthly_sales_labels_json': json.dumps(monthly_sales_labels),
        'monthly_sales_values_json': json.dumps(monthly_sales_values),
        'monthly_growth': round(monthly_growth, 1),
        'comparison_labels_json': json.dumps(comparison_labels),
        'comparison_sales_json': json.dumps(comparison_sales),
        'comparison_expenses_json': json.dumps(comparison_expenses),
        'expense_dates_json': json.dumps(expense_dates),
        'expense_values_json': json.dumps(expense_values),
        'weekday_labels_json': json.dumps(weekday_labels),
        'weekday_values_json': json.dumps(weekday_values),
        'yearly_labels_json': json.dumps(yearly_labels),
        'yearly_values_json': json.dumps(yearly_values),
        
        # Averages for cards
        'daily_avg_sales': avg_sales,
        'monthly_avg_sales': sum(monthly_sales_values) / len(monthly_sales_values) if monthly_sales_values else 0,
        
        # Daily sales trend data (all time)
        'sales_dates_json': json.dumps([s['date'].strftime('%Y-%m-%d') for s in DailySales.objects.all().order_by('date').values('date', 'total_sales')]),
        'sales_values_json': json.dumps([safe_float(s['total_sales']) for s in DailySales.objects.all().order_by('date').values('date', 'total_sales')]),
        
        # Combined tab data
        'net_contribution': total_sales - (total_labor + total_delivery + total_expense),
        'contribution_color': 'success' if (total_sales - (total_labor + total_delivery + total_expense)) >= 0 else 'danger',
    }
    
    return render(request, 'admin/admin_analytics.html', context)


@staff_required
def analytics_api(request):
    """
    Unified Analytics API endpoint for dynamic filtering
    
    Accepts:
        - start_date: YYYY-MM-DD format
        - end_date: YYYY-MM-DD format  
        - month: YYYY-MM format (overrides start_date/end_date if provided)
        - mode: 'sales' | 'expenses' | 'combined'
        
    Returns JSON with:
        - Summary cards data
        - Insight cards (best/worst days, averages)
        - Chart data for the selected period
    """
    from django.db.models import Sum, Avg, Min, Max
    from django.db.models.functions import TruncMonth, TruncYear
    from firstApp.models import DailySales, DailyExpenditure
    from calendar import monthrange
    
    def safe_float(val):
        return float(val) if val else 0.0
    
    # Parse filter parameters
    start_date_str = request.GET.get('start_date', '')
    end_date_str = request.GET.get('end_date', '')
    month_str = request.GET.get('month', '')
    mode = request.GET.get('mode', 'sales')  # sales, expenses, combined
    
    # Determine date range
    start_date = None
    end_date = None
    
    if month_str:
        # Month takes priority - parse YYYY-MM format
        try:
            year, month = map(int, month_str.split('-'))
            start_date = datetime.date(year, month, 1)
            _, last_day = monthrange(year, month)
            end_date = datetime.date(year, month, last_day)
        except (ValueError, TypeError):
            pass
    elif start_date_str and end_date_str:
        # Use custom date range
        try:
            start_date = datetime.datetime.strptime(start_date_str, '%Y-%m-%d').date()
            end_date = datetime.datetime.strptime(end_date_str, '%Y-%m-%d').date()
        except ValueError:
            pass
    
    # Build querysets with filters
    sales_qs = DailySales.objects.all()
    expenses_qs = DailyExpenditure.objects.all()
    
    if start_date and end_date:
        sales_qs = sales_qs.filter(date__range=[start_date, end_date])
        expenses_qs = expenses_qs.filter(date__range=[start_date, end_date])
    
    # Calculate number of days in range for averages
    if start_date and end_date:
        days_in_range = (end_date - start_date).days + 1
    else:
        # Use total days with data
        days_in_range = sales_qs.count() or 1
    
    # Calculate months in range for monthly average
    if start_date and end_date:
        months_diff = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
        months_in_range = max(months_diff, 1)
    else:
        months_in_range = 1
    
    # Prepare response based on mode
    response_data = {
        'success': True,
        'timestamp': timezone.now().isoformat(),
        'mode': mode,
        'filters': {
            'start_date': start_date_str if start_date else '',
            'end_date': end_date_str if end_date else '',
            'month': month_str,
        }
    }
    
    if mode == 'sales':
        # ===== SALES DATA =====
        sales_agg = sales_qs.aggregate(
            total=Sum('total_sales'),
            online=Sum('online_received'),
            cash=Sum('cash_received'),
            labor=Sum('labor_charge'),
            delivery=Sum('delivery_charge'),
            count=Count('id')
        )
        
        total_sales = safe_float(sales_agg['total'])
        total_online = safe_float(sales_agg['online'])
        total_cash = safe_float(sales_agg['cash'])
        total_labor = safe_float(sales_agg['labor'])
        total_delivery = safe_float(sales_agg['delivery'])
        sales_count = sales_agg['count'] or 0
        
        # Calculate percentages
        online_pct = (total_online / total_sales * 100) if total_sales > 0 else 0
        
        # Calculate averages
        daily_avg = (total_sales / days_in_range) if days_in_range > 0 else 0
        monthly_avg = (total_sales / months_in_range) if months_in_range > 0 else 0
        
        # Best/Worst days
        best_day = sales_qs.order_by('-total_sales').first()
        worst_day = sales_qs.filter(total_sales__gt=0).order_by('total_sales').first()
        
        # Best month (within filtered range)
        monthly_data = sales_qs.annotate(
            month=TruncMonth('date')
        ).values('month').annotate(
            total=Sum('total_sales')
        ).order_by('-total')
        
        best_month_data = monthly_data.first() if monthly_data else None
        
        response_data['summary'] = {
            'total_sales': round(total_sales, 2),
            'avg_daily': round(daily_avg, 2),
            'online_amount': round(total_online, 2),
            'online_percentage': round(online_pct, 1),
            'labor_total': round(total_labor, 2),
            'delivery_total': round(total_delivery, 2),
            'cash_total': round(total_cash, 2),
        }
        
        response_data['insights'] = {
            'best_day': {
                'date': best_day.date.strftime('%d %b, %Y') if best_day else 'No data',
                'amount': round(safe_float(best_day.total_sales), 2) if best_day else 0,
            },
            'worst_day': {
                'date': worst_day.date.strftime('%d %b, %Y') if worst_day else 'No data',  
                'amount': round(safe_float(worst_day.total_sales), 2) if worst_day else 0,
            },
            'best_month': {
                'month': best_month_data['month'].strftime('%B %Y') if best_month_data else 'No data',
                'amount': round(safe_float(best_month_data['total']), 2) if best_month_data else 0,
            },
            'averages': {
                'daily': round(daily_avg, 2),
                'monthly': round(monthly_avg, 2),
            }
        }
        
        # Chart data for filtered range
        daily_data = sales_qs.order_by('date').values('date', 'total_sales')
        response_data['chart'] = {
            'dates': [d['date'].strftime('%Y-%m-%d') for d in daily_data],
            'values': [round(safe_float(d['total_sales']), 2) for d in daily_data],
        }
        
    elif mode == 'expenses':
        # ===== EXPENSES DATA =====
        expense_agg = expenses_qs.aggregate(
            total=Sum('total'),
            online=Sum('online_amount'),
            cash=Sum('cash_amount'),
            count=Count('id')
        )
        
        total_expense = safe_float(expense_agg['total'])
        total_online_exp = safe_float(expense_agg['online'])
        total_cash_exp = safe_float(expense_agg['cash'])
        expense_count = expense_agg['count'] or 0
        
        # Calculate percentages
        online_exp_pct = (total_online_exp / total_expense * 100) if total_expense > 0 else 0
        
        # Calculate averages
        daily_avg = (total_expense / days_in_range) if days_in_range > 0 else 0
        monthly_avg = (total_expense / months_in_range) if months_in_range > 0 else 0
        
        # Best/Worst expense days
        highest_exp_day = expenses_qs.order_by('-total').first()
        lowest_exp_day = expenses_qs.filter(total__gt=0).order_by('total').first()
        
        # Highest expense month
        monthly_exp_data = expenses_qs.annotate(
            month=TruncMonth('date')
        ).values('month').annotate(
            total=Sum('total')
        ).order_by('-total')
        
        highest_month_data = monthly_exp_data.first() if monthly_exp_data else None
        
        response_data['summary'] = {
            'total_expenses': round(total_expense, 2),
            'avg_daily': round(daily_avg, 2),
            'online_amount': round(total_online_exp, 2),
            'online_percentage': round(online_exp_pct, 1),
            'cash_total': round(total_cash_exp, 2),
        }
        
        response_data['insights'] = {
            'highest_day': {
                'date': highest_exp_day.date.strftime('%d %b, %Y') if highest_exp_day else 'No data',
                'amount': round(safe_float(highest_exp_day.total), 2) if highest_exp_day else 0,
            },
            'lowest_day': {
                'date': lowest_exp_day.date.strftime('%d %b, %Y') if lowest_exp_day else 'No data',
                'amount': round(safe_float(lowest_exp_day.total), 2) if lowest_exp_day else 0,
            },
            'highest_month': {
                'month': highest_month_data['month'].strftime('%B %Y') if highest_month_data else 'No data',
                'amount': round(safe_float(highest_month_data['total']), 2) if highest_month_data else 0,
            },
            'averages': {
                'daily': round(daily_avg, 2),
                'monthly': round(monthly_avg, 2),
            }
        }
        
        # Chart data for filtered range
        daily_exp_data = expenses_qs.order_by('date').values('date', 'total')
        response_data['chart'] = {
            'dates': [d['date'].strftime('%Y-%m-%d') for d in daily_exp_data],
            'values': [round(safe_float(d['total']), 2) for d in daily_exp_data],
        }
        
    else:  # combined
        # ===== COMBINED DATA (Sales - Expenses) =====
        sales_agg = sales_qs.aggregate(
            total=Sum('total_sales'),
            labor=Sum('labor_charge'),
            delivery=Sum('delivery_charge')
        )
        expense_agg = expenses_qs.aggregate(
            total=Sum('total')
        )
        
        total_sales = safe_float(sales_agg['total'])
        total_labor = safe_float(sales_agg['labor'])
        total_delivery = safe_float(sales_agg['delivery'])
        total_expense = safe_float(expense_agg['total'])
        
        # Net Contribution: Sales - (Labor + Delivery + Expenses)
        net_contribution = total_sales - (total_labor + total_delivery + total_expense)
        contribution_color = 'success' if net_contribution >= 0 else 'danger'
        
        # Calculate averages
        daily_avg = (net_contribution / days_in_range) if days_in_range > 0 else 0
        monthly_avg = (net_contribution / months_in_range) if months_in_range > 0 else 0
        
        # Get monthly comparison data
        sales_monthly = sales_qs.annotate(
            month=TruncMonth('date')
        ).values('month').annotate(
            total=Sum('total_sales')
        ).order_by('month')
        
        expenses_monthly = expenses_qs.annotate(
            month=TruncMonth('date')
        ).values('month').annotate(
            total=Sum('total')
        ).order_by('month')
        
        # Build comparison chart
        expenses_dict = {e['month']: safe_float(e['total']) for e in expenses_monthly}
        comparison_labels = [s['month'].strftime('%b %Y') for s in sales_monthly]
        comparison_sales = [round(safe_float(s['total']), 2) for s in sales_monthly]
        comparison_expenses = [round(expenses_dict.get(s['month'], 0), 2) for s in sales_monthly]
        
        response_data['summary'] = {
            'total_sales': round(total_sales, 2),
            'total_expenses': round(total_expense, 2),
            'labor_total': round(total_labor, 2),
            'delivery_total': round(total_delivery, 2),
            'net_contribution': round(net_contribution, 2),
            'contribution_color': contribution_color,
        }
        
        response_data['insights'] = {
            'averages': {
                'daily': round(daily_avg, 2),
                'monthly': round(monthly_avg, 2),
            }
        }
        
        response_data['chart'] = {
            'labels': comparison_labels,
            'sales': comparison_sales,
            'expenses': comparison_expenses,
        }
    
    return JsonResponse(response_data)

def handle_analytics_export(format_type, data):
    """
    Handle export in PDF, Excel, CSV, or PNG format
    """
    if format_type == 'csv':
        import csv
        from io import StringIO
        
        output = StringIO()
        writer = csv.writer(output)
        writer.writerow(['Period', 'Sales', 'Expenses', 'Net'])
        
        for i, sale in enumerate(data['sales']):
            expense = data['expenses'][i] if i < len(data['expenses']) else {'total': 0}
            writer.writerow([
                sale['month'].strftime('%Y-%m'),
                sale['total'],
                expense['total'],
                sale['total'] - expense['total']
            ])
        
        response = HttpResponse(output.getvalue(), content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="analytics_{timezone.now().strftime("%Y%m%d")}.csv"'
        return response
    
    elif format_type == 'excel':
        try:
            import openpyxl
            from openpyxl.utils import get_column_letter
            from io import BytesIO
            
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Analytics Report"
            
            # Headers
            ws['A1'] = 'Analytics Report'
            ws['A2'] = f'Generated: {timezone.now().strftime("%Y-%m-%d %H:%M")}'
            ws['A3'] = f'Period: {data["start_date"]} to {data["end_date"]}'
            
            # Data headers
            ws['A5'] = 'Month'
            ws['B5'] = 'Sales'
            ws['C5'] = 'Expenses'
            ws['D5'] = 'Net Contribution'
            
            row = 6
            for i, sale in enumerate(data['sales']):
                expense = data['expenses'][i] if i < len(data['expenses']) else {'total': 0}
                ws[f'A{row}'] = sale['month'].strftime('%B %Y')
                ws[f'B{row}'] = float(sale['total'])
                ws[f'C{row}'] = float(expense['total'])
                ws[f'D{row}'] = float(sale['total']) - float(expense['total'])
                row += 1
            
            # Summary
            ws[f'A{row+1}'] = 'TOTAL'
            ws[f'B{row+1}'] = data['total_sales']
            ws[f'C{row+1}'] = data['total_expense']
            ws[f'D{row+1}'] = data['net_contribution']
            
            output = BytesIO()
            wb.save(output)
            output.seek(0)
            
            response = HttpResponse(
                output.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename="analytics_{timezone.now().strftime("%Y%m%d")}.xlsx"'
            return response
        except ImportError:
            # Fallback to CSV if openpyxl not installed
            return handle_analytics_export('csv', data)
    
    else:
        # Default to CSV
        return handle_analytics_export('csv', data)

@staff_member_required
def admin_activity_log_view(request):
    logs = AdminActivityLog.objects.select_related('admin').all().order_by('-timestamp')
    
    # Filters
    action_filter = request.GET.get('action')
    module_filter = request.GET.get('module')
    admin_filter = request.GET.get('admin')
    search_query = request.GET.get('search')
    
    if action_filter:
        logs = logs.filter(action=action_filter)
    if module_filter:
        logs = logs.filter(module=module_filter)
    if admin_filter:
        logs = logs.filter(admin_id=admin_filter)
    if search_query:
        logs = logs.filter(description__icontains=search_query)

    # Pagination
    from django.core.paginator import Paginator
    paginator = Paginator(logs, 20) # Show 20 logs per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Context Data
    from .models import CustomUser
    admins = CustomUser.objects.filter(is_staff=True)
    
    context = {
        'page_obj': page_obj,
        'action_choices': AdminActivityLog.ACTION_CHOICES,
        'module_choices': AdminActivityLog.MODULE_CHOICES,
        'admins': admins,
    }
    return render(request, 'admin/admin_activity_log.html', context)

@staff_member_required
def terminate_session(request, session_id):
    session = get_object_or_404(AdminSession, pk=session_id)
    session.delete()
    messages.success(request, "Session terminated.")
    return redirect('admin_dashboard')

@staff_member_required
def admin_appointment_list(request):
    appointments = Appointment.objects.all()
    return render(request, 'admin/admin_appointments.html', {'appointments': appointments})

@staff_member_required
def admin_appointment_update(request, pk):
    appointment = get_object_or_404(Appointment, pk=pk)
    if request.method == 'POST':
        status = request.POST.get('status')
        extra_charge = request.POST.get('extra_charge')
        electrician_id = request.POST.get('assigned_electrician')
        
        old_status = appointment.status
        old_electrician = appointment.assigned_electrician

        if status:
            appointment.status = status
            
        # Note: visiting_charge is read-only and set during appointment creation
        # It should not be modified from the admin panel
        
        if extra_charge:
            appointment.extra_charge = extra_charge
        
        # Handle electrician assignment
        if electrician_id:
            if electrician_id == 'none':
                appointment.assigned_electrician = None
            else:
                from .models import Electrician
                try:
                    electrician = Electrician.objects.get(id=electrician_id)
                    appointment.assigned_electrician = electrician
                except Electrician.DoesNotExist:
                    pass
            
        appointment.save()
        
        # Send Email Notifications
        try:
            from .email_utils import send_appointment_status_email, send_appointment_complete_email, send_electrician_assignment_email
            
            # If status changed
            if old_status != appointment.status and appointment.email:
                if appointment.status == 'Completed':
                    # Send completion email with review request
                    send_appointment_complete_email(appointment)
                else:
                    # Send status update email (includes electrician info if assigned)
                    send_appointment_status_email(appointment)
            
            # If electrician was newly assigned or changed
            if appointment.assigned_electrician and old_electrician != appointment.assigned_electrician:
                # Send email to user about electrician assignment
                if appointment.email and old_electrician != appointment.assigned_electrician:
                    send_appointment_status_email(appointment)
                
                # Send email to electrician about new appointment
                if appointment.assigned_electrician.email:
                    send_electrician_assignment_email(appointment, appointment.assigned_electrician)
                    
        except Exception as e:
            print(f"Error sending appointment email: {e}")

        messages.success(request, "Appointment updated successfully.")
        return redirect('admin_appointment_list')
    
    # Get all electricians for dropdown
    from .models import Electrician
    electricians = Electrician.objects.filter(is_active=True).order_by('name')
    
    return render(request, 'admin/appointment_update.html', {
        'appointment': appointment,
        'status_choices': Appointment.STATUS_CHOICES,
        'electricians': electricians
    })

@staff_member_required
def admin_appointment_delete(request, pk):
    appointment = get_object_or_404(Appointment, pk=pk)
    appointment.delete()
    messages.success(request, "Appointment deleted successfully.")
    return redirect('admin_appointment_list')



    if not os.path.exists(data_dir):
        os.makedirs(data_dir)

    # Handle File Upload
    if request.method == 'POST' and request.FILES.get('upload_file'):
        uploaded_file = request.FILES['upload_file']
        if uploaded_file.name.endswith(('.xlsx', '.tsv', '.csv')):
            fs = FileSystemStorage(location=data_dir)
            filename = fs.save(uploaded_file.name, uploaded_file)
            messages.success(request, f"File '{filename}' uploaded successfully.")
            return redirect('admin_analytics')
        else:
            messages.error(request, "Invalid file format. Please upload .xlsx, .tsv, or .csv.")
    
    data_files = [f for f in os.listdir(data_dir) if f.endswith(('.xlsx', '.tsv', '.csv'))]
    
    context = {'files': data_files}
    
    selected_file = request.GET.get('file')
    if selected_file and selected_file in data_files:
        file_path = os.path.join(data_dir, selected_file)
        try:
            if selected_file.endswith('.tsv'):
                df = pd.read_csv(file_path, sep='\t')
            elif selected_file.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Basic Analysis Logic (Assuming standard columns: 'Product', 'Quantity', 'Price', 'Delivery', 'Distance')
            # Adjust column names based on actual file structure or requirement
            # If columns don't exist, we just show raw data
            
            analysis = {}
            
            # Normalize column names to lowercase for safer access
            df.columns = df.columns.str.lower()
            
            # 1. Total Revenue
            if 'price' in df.columns:
                analysis['total_revenue'] = df['price'].sum()
            elif 'total' in df.columns:
                 analysis['total_revenue'] = df['total'].sum()
            else:
                 analysis['total_revenue'] = 0
                 
            # 2. Total Delivery Charges
            if 'delivery' in df.columns:
                analysis['total_delivery'] = df['delivery'].sum()
            elif 'delivery_charge' in df.columns:
                analysis['total_delivery'] = df['delivery_charge'].sum()
            else:
                analysis['total_delivery'] = 0
                
            # 3. Total Distance
            if 'distance' in df.columns:
                analysis['total_distance'] = df['distance'].sum()
            else:
                analysis['total_distance'] = 0

            # 4. Product Sales
            if 'product' in df.columns and 'quantity' in df.columns and 'price' in df.columns:
               product_sales = df.groupby('product').agg({'quantity': 'sum', 'price': 'sum'}).reset_index()
               product_sales.columns = ['name', 'quantity', 'sales']
               analysis['product_sales'] = product_sales.to_dict('records')
            else:
               analysis['product_sales'] = []

            # 5. Raw Data HTML
            analysis['html_table'] = df.to_html(classes='table table-bordered table-hover mb-0', index=False)
            
            # 6. Trend Prediction (Linear Regression)
            if 'date' in df.columns:
                target_col = None
                if 'price' in df.columns: target_col = 'price'
                elif 'total' in df.columns: target_col = 'total'
                elif 'sales' in df.columns: target_col = 'sales'
                
                if target_col:
                    try:
                        # Prepare data
                        df_trend = df.copy()
                        df_trend['date'] = pd.to_datetime(df_trend['date'])
                        daily_data = df_trend.groupby('date')[target_col].sum().reset_index().sort_values('date')
                        
                        if len(daily_data) > 1:
                            # Create ordinal dates for regression
                            daily_data['ordinal'] = daily_data['date'].apply(lambda x: x.toordinal())
                            
                            X = daily_data['ordinal'].values
                            y = daily_data[target_col].values
                            
                            # Linear Regression: y = mx + c
                            slope, intercept = np.polyfit(X, y, 1)
                            
                            # Predictions for existing dates (Trend Line)
                            trend_values = slope * X + intercept
                            
                            # Future Predictions (Next 7 days)
                            last_date = daily_data['date'].max()
                            future_dates = [last_date + datetime.timedelta(days=i) for i in range(1, 8)]
                            future_ordinals = np.array([d.toordinal() for d in future_dates])
                            future_predictions = slope * future_ordinals + intercept
                            
                            # Prepare Labels and Data for Chart
                            # Labels: History + Future
                            history_dates_str = [d.strftime('%Y-%m-%d') for d in daily_data['date']]
                            future_dates_str = [d.strftime('%Y-%m-%d') for d in future_dates]
                            all_labels = history_dates_str + future_dates_str
                            
                            actual_data = list(daily_data[target_col]) + [None] * len(future_dates)
                            
                            # Trend Data: History Trend + Future Predictions
                            trend_line_data = list(trend_values) + list(future_predictions)
                            
                            analysis['chart_labels'] = json.dumps(all_labels)
                            analysis['chart_actual'] = json.dumps(actual_data)
                            analysis['chart_trend'] = json.dumps(trend_line_data)
                            analysis['has_prediction'] = True
                            
                    except Exception as reg_err:
                        print(f"Regression error: {reg_err}")
                        
            context['analysis'] = analysis
            context['current_file'] = selected_file
            
        except Exception as e:
            context['error'] = f"Error processing file: {str(e)}"
            
    return render(request, 'admin/admin_analytics.html', context)

@staff_member_required
def admin_delete_analytics_file(request):
    if request.method == 'POST':
        filename = request.POST.get('filename')
        if filename:
            data_dir = os.path.join(settings.MEDIA_ROOT, 'data')
            file_path = os.path.join(data_dir, filename)
            
            # Security check: ensure file is within data_dir
            # Using abspath for safer comparison
            abs_data_dir = os.path.abspath(data_dir)
            abs_file_path = os.path.abspath(file_path)
            
            if abs_file_path.startswith(abs_data_dir) and os.path.exists(abs_file_path):
                try:
                    os.remove(abs_file_path)
                    messages.success(request, f"File '{filename}' deleted successfully.")
                except Exception as e:
                    messages.error(request, f"Error deleting file: {e}")
            else:
                 messages.error(request, "Invalid file path.")
        else:
             messages.error(request, "No filename provided.")
             
    return redirect('admin_analytics')

# ------------------------------------------------------------------
# Product Management
# ------------------------------------------------------------------

@staff_member_required
def admin_product_list(request):
    """
    Admin view to list products with search, sorting, and filtering.
    Enhanced to search across multiple fields and sort by effective price.
    """
    # Base queryset
    products = Product.objects.select_related('category').all()
    
    # Get filter parameters
    search = request.GET.get('search', '').strip()
    visibility = request.GET.get('visibility', '')
    stock = request.GET.get('stock', '')
    sort = request.GET.get('sort', 'newest')
    
    # Apply search filter
    if search:
        products = products.filter(
            Q(name__icontains=search) | 
            Q(brand__icontains=search) |
            Q(category__name__icontains=search) |
            Q(description__icontains=search) |
            Q(vendor__icontains=search)
        ).distinct()
    
    # Apply visibility filter
    if visibility == 'visible':
        products = products.filter(is_visible_on_website=True)
    elif visibility == 'hidden':
        products = products.filter(is_visible_on_website=False)
    
    # Apply stock filter
    if stock == 'in_stock':
        products = products.filter(stock_quantity__gt=0)
    elif stock == 'out_of_stock':
        products = products.filter(stock_quantity=0)
    
    # Annotate with effective price for accurate sorting (handles discounts)
    products = products.annotate(
        effective_price=Coalesce('discount_price', 'price')
    )
    
    # Apply sorting
    if sort == 'name_az':
        products = products.order_by('name')
    elif sort == 'name_za':
        products = products.order_by('-name')
    elif sort == 'price_low':
        products = products.order_by('effective_price')
    elif sort == 'price_high':
        products = products.order_by('-effective_price')
    elif sort == 'stock':
        products = products.order_by('-stock_quantity')
    else:
        # Default to newest products
        products = products.order_by('-created_at')
    
    # Pagination
    paginator = Paginator(products, 25)
    page_number = request.GET.get('page')
    products_page = paginator.get_page(page_number)
    
    return render(request, 'admin/admin_product_list.html', {
        'products': products_page,
        'search': search,
        'visibility': visibility,
        'stock': stock,
        'sort': sort,
    })

@staff_member_required
def admin_product_add(request):
    if request.method == 'POST':
        # Manually handling form for simplicity or use Django Forms
        name = request.POST.get('name')
        category_id = request.POST.get('category')
        description = request.POST.get('description')
        price = request.POST.get('price')
        discount_price = request.POST.get('discount_price') or None
        stock_quantity = request.POST.get('stock_quantity')
        brand = request.POST.get('brand')
        is_trending = request.POST.get('is_trending') == 'on'
        
        main_image = request.FILES.get('image')

        try:
            category = Category.objects.get(id=category_id)
            product = Product.objects.create(
                name=name,
                category=category,
                description=description,
                price=price,
                discount_price=discount_price,
                stock_quantity=stock_quantity,
                brand=brand,
                is_trending=is_trending,
                image=main_image
            )
            
            # Handle multiple images
            files = request.FILES.getlist('more_images')
            for f in files:
                ProductImage.objects.create(product=product, image=f)

            messages.success(request, f"Product '{name}' added successfully.")
            
            # Log Activity
            AdminActivityLog.objects.create(
                admin=request.user,
                action='CREATE',
                module='PRODUCT',
                description=f"Added product: {name}",
                ip_address=request.META.get('REMOTE_ADDR')
            )
            return redirect('admin_product_list')
        except Exception as e:
            messages.error(request, f"Error adding product: {str(e)}")

    categories = Category.objects.all().order_by('name')
    return render(request, 'admin/admin_product_form.html', {'categories': categories})

@staff_member_required
def admin_product_edit(request, pk):
    product = get_object_or_404(Product, pk=pk)
    if request.method == 'POST':
        product.name = request.POST.get('name')
        product.category_id = request.POST.get('category')
        product.description = request.POST.get('description')
        product.price = request.POST.get('price')
        dp = request.POST.get('discount_price')
        product.discount_price = dp if dp else None
        product.stock_quantity = request.POST.get('stock_quantity')
        product.brand = request.POST.get('brand')
        product.is_trending = request.POST.get('is_trending') == 'on'
        
        if request.FILES.get('image'):
            product.image = request.FILES.get('image')
        
        product.save()

        # Handle multiple images (add more)
        files = request.FILES.getlist('more_images')
        for f in files:
            ProductImage.objects.create(product=product, image=f)

        messages.success(request, f"Product '{product.name}' updated successfully.")
        
        AdminActivityLog.objects.create(
            admin=request.user,
            action='UPDATE',
            module='PRODUCT',
            description=f"Updated product: {product.name}",
            ip_address=request.META.get('REMOTE_ADDR')
        )
        return redirect('admin_product_list')

    categories = Category.objects.all().order_by('name')
    return render(request, 'admin/admin_product_form.html', {'product': product, 'categories': categories})

@staff_member_required
def admin_product_delete(request, pk):
    product = get_object_or_404(Product, pk=pk)
    name = product.name
    product.delete()
    messages.success(request, f"Product '{name}' deleted.")
    
    AdminActivityLog.objects.create(
        admin=request.user,
        action='DELETE',
        module='PRODUCT',
        description=f"Deleted product: {name}",
        ip_address=request.META.get('REMOTE_ADDR')
    )
    return redirect('admin_product_list')

# ------------------------------------------------------------------
# Order Management
# ------------------------------------------------------------------

@staff_member_required
def admin_order_list(request):
    orders = Order.objects.select_related('user').all().order_by('-created_at')
    return render(request, 'admin/admin_order_list.html', {'orders': orders})

@staff_member_required
def admin_order_detail(request, pk):
    order = get_object_or_404(Order, pk=pk)
    
    if request.method == 'POST':
        # Handle Update Item Prices (for enquiry-based orders)
        if 'update_item_prices' in request.POST:
            from decimal import Decimal
            
            # Update individual item prices
            subtotal = Decimal('0.00')
            for item in order.items.all():
                price_key = f'item_price_{item.id}'
                if price_key in request.POST:
                    try:
                        new_price = Decimal(request.POST.get(price_key, '0'))
                        item.price = new_price
                        item.save()
                        subtotal += new_price * item.quantity
                    except:
                        pass
            
            # Update delivery charge
            delivery_charge = request.POST.get('delivery_charge', '0')
            try:
                order.delivery_charge = Decimal(delivery_charge)
            except:
                pass
            
            # Update total price
            order.total_price = subtotal
            order.save()
            
            messages.success(request, "Item prices and delivery charge updated successfully.")
            return redirect('admin_order_detail', pk=pk)
        
        # Update Status
        if 'update_status' in request.POST:
            from decimal import Decimal
            
            new_status = request.POST.get('status')
            confirm_pricing = request.POST.get('confirm_pricing') == '1'
            
            # Capture old status to detect change
            old_status = order.status
            
            # Handle pricing confirmation (for enquiry orders)
            if confirm_pricing and order.order_type == 'enquiry' and not order.pricing_confirmed:
                # Mark pricing as confirmed
                order.pricing_confirmed = True
                
                # Calculate final price
                subtotal = sum(item.price * item.quantity for item in order.items.all())
                order.total_price = subtotal
                
                # Ensure delivery_charge is not None
                delivery_charge = order.delivery_charge if order.delivery_charge is not None else Decimal('0.00')
                order.final_price = subtotal + delivery_charge
                order.delivery_charge_status = 'CONFIRMED'
                
                # NOW deduct stock (was deferred until pricing confirmed)
                for item in order.items.all():
                    if item.product.stock_quantity >= item.quantity:
                        item.product.stock_quantity -= item.quantity
                        item.product.save()
                    else:
                        messages.warning(request, f"Insufficient stock for {item.product.name}. Available: {item.product.stock_quantity}")
                
                # Check free delivery eligibility
                if order.user.free_delivery_used_count == 0 and order.distance_km and order.distance_km <= 2:
                    order.free_delivery_applied = True
                    order.delivery_charge = Decimal('0.00')
                    order.final_price = order.total_price
                    order.user.free_delivery_used_count += 1
                    order.user.save()
                    messages.info(request, "Free delivery applied (first order within 2 KM).")
                
                messages.success(request, "Pricing confirmed! Stock has been deducted.")
            
            order.status = new_status
            
            # Handle status-specific actions
            if new_status == 'Confirmed':
                final_price = request.POST.get('final_price')
                delivery_charge = request.POST.get('delivery_charge_status')
                if final_price:
                    try:
                        order.final_price = Decimal(final_price)
                    except:
                        pass
                if delivery_charge:
                    try:
                        order.delivery_charge = Decimal(delivery_charge)
                    except:
                        pass
                
                # Generate receipt number when order is confirmed
                if not order.receipt_number:
                    order.generate_receipt_number()
                    if order.receipt_qr_data and not order.receipt_qr_code:
                        order.generate_receipt_qr_code()
                    messages.success(request, f"Receipt generated: {order.receipt_number}")
            
            # If Out for Delivery, generate OTP
            if new_status == 'Out for Delivery':
                import random
                otp = str(random.randint(100000, 999999))
                order.delivery_otp = otp
                
                # Generate receipt if not already generated
                if not order.receipt_number:
                    order.generate_receipt_number()
                    if order.receipt_qr_data and not order.receipt_qr_code:
                        order.generate_receipt_qr_code()
                
                from .email_utils import send_delivery_otp_email
                send_delivery_otp_email(order, otp)
                
                messages.info(request, f"Delivery OTP generated and sent: {otp}")
            
            # Send status update email for status changes
            if new_status != old_status:
                from .email_utils import send_order_status_email, send_order_delivered_email
                
                if new_status == 'Delivered':
                    # Generate receipt if not already generated
                    if not order.receipt_number:
                        order.generate_receipt_number()
                        if order.receipt_qr_data and not order.receipt_qr_code:
                            order.generate_receipt_qr_code()
                    send_order_delivered_email(order)
                elif new_status == 'Out for Delivery':
                    # Already sent OTP email above
                    pass
                elif new_status == 'Cancelled':
                    send_order_status_email(order)
                elif new_status in ['Confirmed', 'Price Shared']:
                    # Send confirmation/price email
                    send_order_status_email(order)

            order.save()
            messages.success(request, f"Order status updated to {new_status}")
            
        return redirect('admin_order_detail', pk=pk)

    return render(request, 'admin/admin_order_detail.html', {'order': order})

# ------------------------------------------------------------------
# Review Management
# ------------------------------------------------------------------

@staff_member_required
def admin_reviews_list(request):
    reviews_list = Review.objects.select_related('product', 'user').all().order_by('-created_at')
    
    paginator = Paginator(reviews_list, 20)
    page_number = request.GET.get('page')
    reviews = paginator.get_page(page_number)
    
    return render(request, 'admin/admin_reviews_list.html', {'reviews': reviews})

@staff_member_required
def admin_delete_review(request, review_id):
    review = get_object_or_404(Review, id=review_id)
    review.delete()
    messages.success(request, "Review deleted successfully.")
    return redirect('admin_reviews_list')

@staff_member_required
def admin_approve_review(request, review_id):
    review = get_object_or_404(Review, id=review_id)
    review.is_approved = not review.is_approved
    review.save()
    status = "visible" if review.is_approved else "hidden"
    messages.success(request, f"Review is now {status}.")
    return redirect('admin_reviews_list')

# ------------------------------------------------------------------
# Daily Sales Management
# ------------------------------------------------------------------

@staff_member_required
def admin_daily_sales(request):
    # Default: Date DESC (most recent first)
    sort_by = request.GET.get('sort_by', 'date')
    sort_order = request.GET.get('sort_order', 'desc')
    
    # Validate sortable fields
    sortable_fields = ['date', 'day', 'total_sales', 'online_received', 'cash_received', 
                       'labor_charge', 'delivery_charge', 'subtotal', 'admin__username']
    
    if sort_by not in sortable_fields:
        sort_by = 'date'
    
    # Build order_by string
    if sort_order == 'desc':
        order_field = f'-{sort_by}'
    else:
        order_field = sort_by
    
    sales_list = DailySales.objects.all().order_by(order_field)
    
    # Filters
    month = request.GET.get('month') # standard format YYYY-MM
    date = request.GET.get('date')
    year = request.GET.get('year')
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    day_name = request.GET.get('day_name')
    
    if month:
        sales_list = sales_list.filter(date__month=month.split('-')[1], date__year=month.split('-')[0])
    if date:
        sales_list = sales_list.filter(date=date)
    if year:
        sales_list = sales_list.filter(date__year=year)
    if start_date and end_date:
        sales_list = sales_list.filter(date__range=[start_date, end_date])
    if day_name:
        sales_list = sales_list.filter(day__iexact=day_name.strip())
    
    # Calculate Totals
    total_sales = sales_list.aggregate(Sum('total_sales'))['total_sales__sum'] or 0
    total_online = sales_list.aggregate(Sum('online_received'))['online_received__sum'] or 0
    total_cash = sales_list.aggregate(Sum('cash_received'))['cash_received__sum'] or 0
    total_labor = sales_list.aggregate(Sum('labor_charge'))['labor_charge__sum'] or 0
    total_delivery = sales_list.aggregate(Sum('delivery_charge'))['delivery_charge__sum'] or 0

    paginator = Paginator(sales_list, 50)  # Show 50 records per page
    page_number = request.GET.get('page')
    sales_records = paginator.get_page(page_number)

    context = {
        'sales_records': sales_records,
        'total_sales': total_sales,
        'total_online': total_online,
        'total_cash': total_cash,
        'total_labor': total_labor,
        'total_delivery': total_delivery,
        'current_sort_by': sort_by,
        'current_sort_order': sort_order,
    }

    return render(request, 'admin/admin_daily_sales.html', context)

@staff_member_required
def admin_add_daily_sales(request):
    if request.method == 'POST':
        form = DailySalesForm(request.POST)
        if form.is_valid():
            sale = form.save(commit=False)
            sale.admin = request.user
            sale.save()
            messages.success(request, "Daily sales record added.")
            return redirect('admin_daily_sales')
    else:
        # Pre-fill date and day
        initial_data = {
            'date': timezone.now().date(),
            'day': timezone.now().strftime('%A')
        }
        form = DailySalesForm(initial=initial_data)
    
    return render(request, 'admin/admin_add_sales.html', {'form': form, 'title': 'Add Sales Entry'})

@staff_member_required
def admin_edit_daily_sales(request, pk):
    sale = get_object_or_404(DailySales, pk=pk)
    if request.method == 'POST':
        form = DailySalesForm(request.POST, instance=sale)
        if form.is_valid():
            form.save()
            messages.success(request, "Daily sales record updated.")
            return redirect('admin_daily_sales')
    else:
        form = DailySalesForm(instance=sale)
    return render(request, 'admin/admin_add_sales.html', {'form': form, 'title': 'Edit Sales Entry'})

@staff_member_required
def admin_delete_daily_sales(request, pk):
    sale = get_object_or_404(DailySales, pk=pk)
    sale.delete()
    messages.success(request, "Sales record deleted.")
    return redirect('admin_daily_sales')

@staff_member_required
def admin_export_sales(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="daily_sales.csv"'
    
    writer = csv.writer(response)
    writer.writerow(['Date', 'Day', 'Total Sales', 'Online Received', 'Cash Received', 'Subtotal', 'Remark', 'Admin', 'Timestamp'])
    
    sales = DailySales.objects.all().order_by('-date')
    for sale in sales:
        writer.writerow([sale.date, sale.day, sale.total_sales, sale.online_received, sale.cash_received, sale.subtotal, sale.remark, sale.admin.username if sale.admin else 'Unknown', sale.created_at])
        
    return response

# ------------------------------------------------------------------
# Daily Expenditure Management
# ------------------------------------------------------------------

@staff_member_required
def admin_daily_expenses(request):
    # Default: Date DESC (most recent first)
    sort_by = request.GET.get('sort_by', 'date')
    sort_order = request.GET.get('sort_order', 'desc')
    
    # Validate sortable fields
    sortable_fields = ['date', 'day', 'online_amount', 'cash_amount', 'total', 'admin__username']
    
    if sort_by not in sortable_fields:
        sort_by = 'date'
    
    # Build order_by string
    if sort_order == 'desc':
        order_field = f'-{sort_by}'
    else:
        order_field = sort_by
    
    expenses_list = DailyExpenditure.objects.all().order_by(order_field, '-created_at')
    
    # Filters
    month = request.GET.get('month')
    date = request.GET.get('date')
    year = request.GET.get('year')
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    day_name = request.GET.get('day_name')
    
    if month:
        expenses_list = expenses_list.filter(date__month=month.split('-')[1], date__year=month.split('-')[0])
    if date:
        expenses_list = expenses_list.filter(date=date)
    if year:
        expenses_list = expenses_list.filter(date__year=year)
    if start_date and end_date:
        expenses_list = expenses_list.filter(date__range=[start_date, end_date])
    if day_name:
        expenses_list = expenses_list.filter(day__iexact=day_name.strip())


    # Calculate Totals (using combined online + cash model)
    total_expenses = expenses_list.aggregate(Sum('total'))['total__sum'] or 0
    total_online_expenses = expenses_list.aggregate(Sum('online_amount'))['online_amount__sum'] or 0
    total_cash_expenses = expenses_list.aggregate(Sum('cash_amount'))['cash_amount__sum'] or 0

    paginator = Paginator(expenses_list, 50)  # Show 50 records per page
    page_number = request.GET.get('page')
    expenses = paginator.get_page(page_number)

    context = {
        'expenses': expenses,
        'total_expenses': total_expenses,
        'total_online_expenses': total_online_expenses,
        'total_cash_expenses': total_cash_expenses,
        'current_sort_by': sort_by,
        'current_sort_order': sort_order,
    }

    return render(request, 'admin/admin_daily_expenses.html', context)

@staff_member_required
def admin_add_daily_expense(request):
    if request.method == 'POST':
        form = DailyExpenditureForm(request.POST)
        if form.is_valid():
            expense = form.save(commit=False)
            expense.admin = request.user
            expense.save()
            messages.success(request, "Daily expenditure record added.")
            return redirect('admin_daily_expenses')
    else:
        initial_data = {
            'date': timezone.now().date(),
        }
        form = DailyExpenditureForm(initial=initial_data)
    
    return render(request, 'admin/admin_add_expense.html', {'form': form, 'title': 'Add Expense'})

@staff_member_required
def admin_edit_daily_expense(request, pk):
    expense = get_object_or_404(DailyExpenditure, pk=pk)
    if request.method == 'POST':
        form = DailyExpenditureForm(request.POST, instance=expense)
        if form.is_valid():
            form.save()
            messages.success(request, "Daily expenditure record updated.")
            return redirect('admin_daily_expenses')
    else:
        form = DailyExpenditureForm(instance=expense)
    return render(request, 'admin/admin_add_expense.html', {'form': form, 'title': 'Edit Expense'})

@staff_member_required
def admin_delete_daily_expense(request, pk):
    expense = get_object_or_404(DailyExpenditure, pk=pk)
    expense.delete()
    messages.success(request, "Expense record deleted.")
    return redirect('admin_daily_expenses')

@staff_member_required
def admin_export_sales(request):
    fmt = request.GET.get('format', 'csv')
    sales = DailySales.objects.all().order_by('-date')
    
    # Define columns - updated to include labor and delivery charges
    columns = ['Date', 'Day', 'Total Sales', 'Online', 'Cash', 'Labor Charge', 'Delivery Charge', 'Subtotal', 'Remark', 'Admin']
    data = []
    for sale in sales:
        data.append([
            str(sale.date), 
            sale.day, 
            str(sale.total_sales), 
            str(sale.online_received), 
            str(sale.cash_received),
            str(sale.labor_charge) if sale.labor_charge is not None else '0.00',
            str(sale.delivery_charge) if sale.delivery_charge is not None else '0.00',
            str(sale.subtotal),
            sale.remark, 
            sale.admin.username if sale.admin else 'Unknown'
        ])

    if fmt == 'tsv':
        response = HttpResponse(content_type='text/tab-separated-values')
        response['Content-Disposition'] = 'attachment; filename="daily_sales.tsv"'
        writer = csv.writer(response, delimiter='\t')
        writer.writerow(columns)
        for row in data:
            writer.writerow(row)
        return response

    elif fmt == 'pdf':
        from reportlab.lib.pagesizes import letter, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet
        
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="daily_sales.pdf"'
        
        # Use landscape mode for wider table with more columns
        doc = SimpleDocTemplate(response, pagesize=landscape(letter))
        elements = []
        styles = getSampleStyleSheet()
        
        elements.append(Paragraph("Daily Sales Report", styles['Title']))
        
        table_data = [columns] + data
        t = Table(table_data)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 0), (-1, -1), 6),
        ]))
        elements.append(t)
        doc.build(elements)
        return response

    elif fmt == 'word':
        from docx import Document
        
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="daily_sales.docx"'
        
        doc = Document()
        doc.add_heading('Daily Sales Report', 0)
        
        table = doc.add_table(rows=1, cols=len(columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr_cells[i].text = col
            
        for row in data:
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
                
        doc.save(response)
        return response

    else: # Default CSV
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="daily_sales.csv"'
        writer = csv.writer(response)
        writer.writerow(columns)
        for row in data:
            writer.writerow(row)
        return response

@staff_member_required
def admin_export_expenses(request):
    fmt = request.GET.get('format', 'csv')
    expenses = DailyExpenditure.objects.all().order_by('-date')
    
    columns = ['Date', 'Day', 'Online Amount', 'Cash Amount', 'Total', 'Description', 'Admin']
    data = []
    for ex in expenses:
        data.append([
            str(ex.date), 
            ex.day, 
            str(ex.online_amount) if ex.online_amount is not None else '0.00',
            str(ex.cash_amount) if ex.cash_amount is not None else '0.00',
            str(ex.total),
            ex.description, 
            ex.admin.username if ex.admin else 'Unknown'
        ])

    if fmt == 'tsv':
        response = HttpResponse(content_type='text/tab-separated-values')
        response['Content-Disposition'] = 'attachment; filename="daily_expenses.tsv"'
        writer = csv.writer(response, delimiter='\t')
        writer.writerow(columns)
        for row in data:
            writer.writerow(row)
        return response
        
    elif fmt == 'pdf':
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="daily_expenses.pdf"'
        
        doc = SimpleDocTemplate(response, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()
        elements.append(Paragraph("Daily Expenses Report", styles['Title']))
        
        table_data = [columns] + data
        t = Table(table_data)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
        ]))
        elements.append(t)
        doc.build(elements)
        return response

    elif fmt == 'word':
        from docx import Document
        
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="daily_expenses.docx"'
        
        doc = Document()
        doc.add_heading('Daily Expenses Report', 0)
        
        table = doc.add_table(rows=1, cols=len(columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr_cells[i].text = col
            
        for row in data:
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
                
        doc.save(response)
        return response

    else: # Default CSV
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="daily_expenses.csv"'
        writer = csv.writer(response)
        writer.writerow(columns)
        for row in data:
            writer.writerow(row)
        return response

# ------------------------------------------------------------------
# Category Management
# ------------------------------------------------------------------

@staff_member_required
def admin_category_list(request):
    categories = Category.objects.annotate(product_count=Count('products')).all()
    return render(request, 'admin/admin_category_list.html', {'categories': categories})

@staff_member_required
def admin_category_add(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        image = request.FILES.get('image')
        
        if name:
            Category.objects.create(name=name, image=image)
            messages.success(request, "Category created successfully.")
            return redirect('admin_category_list')
        
    return render(request, 'admin/admin_category_form.html')

@staff_member_required
def admin_category_edit(request, pk):
    category = get_object_or_404(Category, pk=pk)
    if request.method == 'POST':
        category.name = request.POST.get('name')
        if request.FILES.get('image'):
            category.image = request.FILES.get('image')
        category.save()
        messages.success(request, "Category updated successfully.")
        return redirect('admin_category_list')
        
    return render(request, 'admin/admin_category_form.html', {'category': category})

@staff_member_required
def admin_category_delete(request, pk):
    category = get_object_or_404(Category, pk=pk)
    category.delete()
    messages.success(request, "Category deleted.")
    return redirect('admin_category_list')

# ------------------------------------------------------------------
# File Upload & Bulk Operations
# ------------------------------------------------------------------

def validate_columns(df, required):
    # Normalize columns: strip, lower, replace space with underscore
    df.columns = df.columns.astype(str).str.strip().str.lower().str.replace(' ', '_')
    missing = [c for c in required if c not in df.columns]
    return missing

@staff_member_required
def admin_upload_sales(request):
    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']
        try:
            if file.name.endswith('.csv'):
                df = pd.read_csv(file)
            elif file.name.endswith('.tsv'):
                df = pd.read_csv(file, sep='\t')
            elif file.name.endswith('.xlsx'):
                df = pd.read_excel(file)
            else:
                messages.error(request, "Invalid file format. Please upload CSV, TSV, or XLSX.")
                return redirect('admin_daily_sales')

            required = ['date', 'total_sales']
            missing = validate_columns(df, required)
            if missing:
                messages.error(request, f"Missing required columns: {', '.join(missing)}")
                return redirect('admin_daily_sales')
                
            count = 0
            for _, row in df.iterrows():
                try:
                    date_val = pd.to_datetime(row['date']).date()
                    DailySales.objects.update_or_create(
                        date=date_val,
                        defaults={
                            'total_sales': row.get('total_sales', 0),
                            'online_received': row.get('online_received', 0),
                            'cash_received': row.get('cash_received', 0),
                            'remark': row.get('remark', 'Updated via Upload'),
                            'admin': request.user
                        }
                    )
                    count += 1
                except Exception as row_err:
                    print(f"Skipping row: {row_err}")

            messages.success(request, f"Successfully processed {count} sales records.")
            return redirect('admin_daily_sales')
            
        except Exception as e:
            messages.error(request, f"Error processing file: {e}")
            return redirect('admin_daily_sales')
            
    return render(request, 'admin/admin_upload_form.html', {'title': 'Upload Daily Sales (CSV/TSV/XLSX)'})

@staff_member_required
def admin_upload_expenses(request):
    if request.method == 'POST' and request.FILES.get('file'):
        file = request.FILES['file']
        try:
            if file.name.endswith('.csv'):
                df = pd.read_csv(file)
            elif file.name.endswith('.tsv'):
                df = pd.read_csv(file, sep='\t')
            elif file.name.endswith('.xlsx'):
                df = pd.read_excel(file)
            else:
                messages.error(request, "Invalid file format. Please upload CSV, TSV, or XLSX.")
                return redirect('admin_daily_expenses')

            # Required: Date, Online Amount, Cash Amount
            # Mappable column names
            required = ['date']
            missing = validate_columns(df, required)
            if missing:
                messages.error(request, f"Missing required columns: {', '.join(missing)}")
                return redirect('admin_daily_expenses')
                
            count = 0
            for _, row in df.iterrows():
                try:
                    date_val = pd.to_datetime(row['date']).date()
                    DailyExpenditure.objects.create(
                        date=date_val,
                        online_amount=row.get('online_amount', 0) or 0,
                        cash_amount=row.get('cash_amount', 0) or 0,
                        description=row.get('description', 'Imported Expense'),
                        admin=request.user
                    )
                    count += 1
                except Exception as row_err:
                    print(f"Skipping row: {row_err}")
                    
            messages.success(request, f"Successfully imported {count} expense records.")
            return redirect('admin_daily_expenses')
            
        except Exception as e:
            messages.error(request, f"Error processing file: {e}")
            return redirect('admin_daily_expenses')
            
    return render(request, 'admin/admin_upload_form.html', {'title': 'Upload Daily Expenses (CSV/TSV/XLSX)'})




# ------------------------------------------------------------------
# User Management
# ------------------------------------------------------------------

@staff_member_required
def admin_user_list(request):
    # Filter for regular users (not superusers or staff) if desired, or all
    # Let's show all users but maybe visually distinguish
    users_list = CustomUser.objects.filter(is_superuser=False, is_staff=False).order_by('-date_joined')
    
    paginator = Paginator(users_list, 20)
    page_number = request.GET.get('page')
    users = paginator.get_page(page_number)
    
    return render(request, 'admin/admin_user_list.html', {'users': users})

@staff_member_required
def admin_user_detail(request, user_id):
    from .models import CustomUser 
    user = get_object_or_404(CustomUser, id=user_id)
    
    # Cart
    try:
        cart_items = user.cart.items.all()
        cart_total = user.cart.total_price
    except:
        cart_items = []
        cart_total = 0
        
    # Wishlist
    # Wishlist is ForeignKey in models.py: user = models.ForeignKey(..., related_name='wishlist')
    wishlist_items = user.wishlist.all()
    
    # Orders (Optional but useful)
    orders = user.orders.all().order_by('-created_at')[:5]
    
    context = {
        'customer': user,
        'cart_items': cart_items,
        'cart_total': cart_total,
        'wishlist_items': wishlist_items,
        'recent_orders': orders
    }
    
    return render(request, 'admin/admin_user_detail.html', context)


# ------------------------------------------------------------------
# TASK 2: Notification Management
# ------------------------------------------------------------------

@staff_member_required
def admin_notifications_list(request):
    """List all notifications"""
    notifications_list = Notification.objects.all().order_by('-created_at')
    
    paginator = Paginator(notifications_list, 20)
    page_number = request.GET.get('page')
    notifications = paginator.get_page(page_number)
    
    return render(request, 'admin/admin_notifications_list.html', {'notifications': notifications})


@staff_member_required
def admin_notification_create(request):
    """Create a new notification"""
    from .models import Notification, UserNotification, CustomUser
    from .forms_notification import NotificationForm
    
    if request.method == 'POST':
        form = NotificationForm(request.POST, request.FILES)
        if form.is_valid():
            notification = form.save(commit=False)
            notification.created_by = request.user
            notification.save()
            form.save_m2m()  # Save many-to-many relationship
            
            # Create UserNotification records for targeted users
            target_users = notification.get_target_users_queryset()
            created_count = 0
            
            for user in target_users:
                UserNotification.objects.get_or_create(
                    user=user,
                    notification=notification
                )
                created_count += 1
            
            messages.success(request, f"Notification created and sent to {created_count} user(s)!")
            
            # Log Activity
            AdminActivityLog.objects.create(
                admin=request.user,
                action='CREATE',
                module='NOTIFICATION',
                description=f"Created notification: {notification.title}",
                ip_address=request.META.get('REMOTE_ADDR')
            )
            
            return redirect('admin_notifications_list')
    else:
        form = NotificationForm()
    
    return render(request, 'admin/admin_notification_form.html', {'form': form, 'title': 'Create Notification'})


@staff_member_required
def admin_notification_edit(request, pk):
    """Edit an existing notification"""
    from .models import Notification
    from .forms_notification import NotificationForm
    
    notification = get_object_or_404(Notification, pk=pk)
    
    if request.method == 'POST':
        form = NotificationForm(request.POST, request.FILES, instance=notification)
        if form.is_valid():
            form.save()
            messages.success(request, "Notification updated successfully!")
            
            AdminActivityLog.objects.create(
                admin=request.user,
                action='UPDATE',
                module='NOTIFICATION',
                description=f"Updated notification: {notification.title}",
                ip_address=request.META.get('REMOTE_ADDR')
            )
            
            return redirect('admin_notifications_list')
    else:
        form = NotificationForm(instance=notification)
    
    return render(request, 'admin/admin_notification_form.html', {
        'form': form,
        'notification': notification,
        'title': 'Edit Notification'
    })


@staff_member_required
def admin_notification_delete(request, pk):
    """Delete a notification"""
    from .models import Notification
    
    notification = get_object_or_404(Notification, pk=pk)
    title = notification.title
    notification.delete()
    
    messages.success(request, f"Notification '{title}' deleted successfully.")
    
    AdminActivityLog.objects.create(
        admin=request.user,
        action='DELETE',
        module='NOTIFICATION',
        description=f"Deleted notification: {title}",
        ip_address=request.META.get('REMOTE_ADDR')
    )
    
    return redirect('admin_notifications_list')


@staff_member_required
def admin_notification_detail(request, pk):
    """View notification details and target audience"""
    from .models import Notification
    
    notification = get_object_or_404(Notification, pk=pk)
    user_notifications = notification.user_notifications.select_related('user').all()[:50]
    total_recipients = notification.user_notifications.count()
    read_count = notification.user_notifications.filter(is_read=True).count()
    
    # Calculate read percentage
    read_percentage = 0
    if total_recipients > 0:
        read_percentage = round((read_count / total_recipients) * 100, 1)
    
    context = {
        'notification': notification,
        'user_notifications': user_notifications,
        'total_recipients': total_recipients,
        'read_count': read_count,
        'unread_count': total_recipients - read_count,
        'read_percentage': read_percentage,
    }
    

    return render(request, 'admin/admin_notification_detail.html', context)


# ------------------------------------------------------------------
# Service Pricing Management
# ------------------------------------------------------------------

@staff_member_required
def admin_service_prices(request):
    """List and manage service prices by service type and zone."""
    prices = ServicePrice.objects.all().order_by('service_type', 'zone')
    
    # Group by service type for better display
    services = {}
    for price in prices:
        if price.service_type not in services:
            services[price.service_type] = []
        services[price.service_type].append(price)
    
    # Get available service types and zones for the form
    service_types = ServicePrice.SERVICE_TYPE_CHOICES
    zones = ServicePrice.ZONE_CHOICES
    
    context = {
        'prices': prices,
        'services': services,
        'service_types': service_types,
        'zones': zones,
    }
    return render(request, 'admin/admin_service_prices.html', context)


@staff_member_required
def admin_service_price_save(request):
    """Save or update a service price."""
    if request.method == 'POST':
        service_type = request.POST.get('service_type')
        zone = request.POST.get('zone')
        base_price = request.POST.get('base_price', 200)
        min_service_charge = request.POST.get('min_service_charge', 300)
        max_service_charge = request.POST.get('max_service_charge', 1500)
        is_active = request.POST.get('is_active') == 'on'
        
        try:
            # Try to get existing record or create new
            price, created = ServicePrice.objects.update_or_create(
                service_type=service_type,
                zone=zone,
                defaults={
                    'base_price': base_price,
                    'min_service_charge': min_service_charge,
                    'max_service_charge': max_service_charge,
                    'is_active': is_active,
                }
            )
            
            action = 'Created' if created else 'Updated'
            messages.success(request, f'{action} pricing for {service_type} in {price.get_zone_display()}')
            
            AdminActivityLog.objects.create(
                admin=request.user,
                action='CREATE' if created else 'UPDATE',
                module='SERVICE_PRICE',
                description=f'{action} {service_type} - {zone}: ₹{base_price}',
                ip_address=request.META.get('REMOTE_ADDR')
            )
        except Exception as e:
            messages.error(request, f'Error saving price: {str(e)}')
    
    return redirect('admin_service_prices')


@staff_member_required
def admin_service_price_delete(request, pk):
    """Delete a service price."""
    price = get_object_or_404(ServicePrice, pk=pk)
    service_info = f"{price.service_type} - {price.get_zone_display()}"
    price.delete()
    
    messages.success(request, f'Deleted pricing for {service_info}')
    
    AdminActivityLog.objects.create(
        admin=request.user,
        action='DELETE',
        module='SERVICE_PRICE',
        description=f'Deleted {service_info}',
        ip_address=request.META.get('REMOTE_ADDR')
    )
    
    return redirect('admin_service_prices')


@staff_member_required
def admin_bulk_create_prices(request):
    """Bulk create default prices for all service-zone combinations."""
    if request.method == 'POST':
        default_base = float(request.POST.get('default_base', 200))
        default_min = float(request.POST.get('default_min', 300))
        default_max = float(request.POST.get('default_max', 1500))
        
        created_count = 0
        for service_type, _ in ServicePrice.SERVICE_TYPE_CHOICES:
            for zone, _ in ServicePrice.ZONE_CHOICES:
                _, created = ServicePrice.objects.get_or_create(
                    service_type=service_type,
                    zone=zone,
                    defaults={
                        'base_price': default_base,
                        'min_service_charge': default_min,
                        'max_service_charge': default_max,
                        'is_active': True,
                    }
                )
                if created:
                    created_count += 1
        
        messages.success(request, f'Created {created_count} new price entries')
        
        AdminActivityLog.objects.create(
            admin=request.user,
            action='CREATE',
            module='SERVICE_PRICE',
            description=f'Bulk created {created_count} price entries',
            ip_address=request.META.get('REMOTE_ADDR')
        )
    
    return redirect('admin_service_prices')


# ------------------------------------------------------------------
# Service Type Management (Dynamic Services)
# ------------------------------------------------------------------

@staff_member_required
def admin_service_types(request):
    """List and manage service types with distance-based pricing."""
    from .models import ServiceType
    
    service_types = ServiceType.objects.all().order_by('display_order', 'name')
    
    context = {
        'service_types': service_types,
        'total_active': service_types.filter(is_active=True).count(),
        'total_inactive': service_types.filter(is_active=False).count(),
        'total_fixed': service_types.filter(pricing_mode='fixed').count(),
        'total_confirm': service_types.filter(pricing_mode='confirm').count(),
    }
    return render(request, 'admin/admin_service_types.html', context)


@staff_member_required
def admin_service_type_save(request):
    """Add or update a service type with distance-based pricing."""
    from .models import ServiceType
    from decimal import Decimal
    
    if request.method == 'POST':
        service_id = request.POST.get('id')
        name = request.POST.get('name', '').strip()
        description = request.POST.get('description', '').strip()
        icon = request.POST.get('icon', 'fa-tools').strip()
        display_order = int(request.POST.get('display_order', 0))
        is_active = request.POST.get('is_active') == 'on'
        pricing_mode = request.POST.get('pricing_mode', 'fixed')
        
        # Parse pricing fields
        default_charge = Decimal(request.POST.get('default_charge', '300'))
        
        def parse_decimal(value):
            if value and value.strip():
                return Decimal(value.strip())
            return None
        
        charge_within_500m = parse_decimal(request.POST.get('charge_within_500m'))
        charge_within_1km = parse_decimal(request.POST.get('charge_within_1km'))
        charge_within_3km = parse_decimal(request.POST.get('charge_within_3km'))
        charge_within_5km = parse_decimal(request.POST.get('charge_within_5km'))
        charge_within_7km = parse_decimal(request.POST.get('charge_within_7km'))
        
        if not name:
            messages.error(request, 'Service name is required.')
            return redirect('admin_service_types')
        
        try:
            if service_id:
                # Update existing
                service = ServiceType.objects.get(id=service_id)
                service.name = name
                service.description = description
                service.icon = icon
                service.display_order = display_order
                service.is_active = is_active
                service.pricing_mode = pricing_mode
                service.default_charge = default_charge
                service.charge_within_500m = charge_within_500m
                service.charge_within_1km = charge_within_1km
                service.charge_within_3km = charge_within_3km
                service.charge_within_5km = charge_within_5km
                service.charge_within_7km = charge_within_7km
                service.save()
                messages.success(request, f'Service "{name}" updated successfully.')
            else:
                # Create new
                ServiceType.objects.create(
                    name=name,
                    description=description,
                    icon=icon,
                    display_order=display_order,
                    is_active=is_active,
                    pricing_mode=pricing_mode,
                    default_charge=default_charge,
                    charge_within_500m=charge_within_500m,
                    charge_within_1km=charge_within_1km,
                    charge_within_3km=charge_within_3km,
                    charge_within_5km=charge_within_5km,
                    charge_within_7km=charge_within_7km,
                )
                messages.success(request, f'Service "{name}" created successfully.')
        except Exception as e:
            messages.error(request, f'Error: {str(e)}')
    
    return redirect('admin_service_types')


@staff_member_required
def admin_service_type_delete(request, pk):
    """Delete a service type."""
    from .models import ServiceType
    
    try:
        service = ServiceType.objects.get(id=pk)
        name = service.name
        service.delete()
        messages.success(request, f'Service "{name}" deleted successfully.')
    except ServiceType.DoesNotExist:
        messages.error(request, 'Service type not found.')
    
    return redirect('admin_service_types')


# ------------------------------------------------------------------
# Distance Zone Management (Distance-Based Pricing)
# ------------------------------------------------------------------

@staff_member_required
def admin_distance_zones(request):
    """List and manage distance zones for pricing."""
    from .models import DistanceZone
    
    zones = DistanceZone.objects.all().order_by('min_distance_km')
    zone_choices = DistanceZone.ZONE_CODES
    
    context = {
        'zones': zones,
        'zone_choices': zone_choices,
    }
    return render(request, 'admin/admin_distance_zones.html', context)


@staff_member_required
def admin_distance_zone_save(request):
    """Add or update a distance zone."""
    from .models import DistanceZone
    
    if request.method == 'POST':
        zone_id = request.POST.get('id')
        code = request.POST.get('code', '').strip()
        min_distance = request.POST.get('min_distance_km', 0)
        max_distance = request.POST.get('max_distance_km', 999)
        base_charge = request.POST.get('base_charge', 200)
        is_active = request.POST.get('is_active') == 'on'
        requires_confirmation = request.POST.get('requires_confirmation') == 'on'
        notes = request.POST.get('notes', '').strip()
        
        try:
            if zone_id:
                # Update existing
                zone = DistanceZone.objects.get(id=zone_id)
                zone.code = code
                zone.min_distance_km = min_distance
                zone.max_distance_km = max_distance
                zone.base_charge = base_charge
                zone.is_active = is_active
                zone.requires_confirmation = requires_confirmation
                zone.notes = notes
                zone.save()
                messages.success(request, f'Distance zone updated successfully.')
            else:
                # Create new
                DistanceZone.objects.create(
                    code=code,
                    min_distance_km=min_distance,
                    max_distance_km=max_distance,
                    base_charge=base_charge,
                    is_active=is_active,
                    requires_confirmation=requires_confirmation,
                    notes=notes
                )
                messages.success(request, f'Distance zone created successfully.')
        except Exception as e:
            messages.error(request, f'Error: {str(e)}')
    
    return redirect('admin_distance_zones')


@staff_member_required
def admin_distance_zone_delete(request, pk):
    """Delete a distance zone."""
    from .models import DistanceZone
    
    try:
        zone = DistanceZone.objects.get(id=pk)
        zone.delete()
        messages.success(request, 'Distance zone deleted successfully.')
    except DistanceZone.DoesNotExist:
        messages.error(request, 'Distance zone not found.')
    
    return redirect('admin_distance_zones')


@staff_member_required
def admin_setup_default_distance_zones(request):
    """Create default distance zones."""
    from .models import DistanceZone
    
    if request.method == 'POST':
        default_zones = [
            {'code': 'VERY_NEAR', 'min': 0, 'max': 0.5, 'charge': 150, 'confirm': False, 'notes': ''},
            {'code': 'NEAR', 'min': 0.5, 'max': 3, 'charge': 200, 'confirm': False, 'notes': ''},
            {'code': 'MEDIUM', 'min': 3, 'max': 5, 'charge': 250, 'confirm': False, 'notes': ''},
            {'code': 'FAR', 'min': 5, 'max': 7, 'charge': 300, 'confirm': False, 'notes': ''},
            {'code': 'VERY_FAR', 'min': 7, 'max': 999, 'charge': 0, 'confirm': True, 'notes': 'Price will be confirmed by staff'},
        ]
        
        created_count = 0
        for zone_data in default_zones:
            _, created = DistanceZone.objects.get_or_create(
                code=zone_data['code'],
                defaults={
                    'min_distance_km': zone_data['min'],
                    'max_distance_km': zone_data['max'],
                    'base_charge': zone_data['charge'],
                    'requires_confirmation': zone_data['confirm'],
                    'notes': zone_data['notes'],
                    'is_active': True,
                }
            )
            if created:
                created_count += 1
        
        messages.success(request, f'Created {created_count} default distance zones.')
    
    return redirect('admin_distance_zones')


# ------------------------------------------------------------------
# Area Region Management (Cities in Regions)
# ------------------------------------------------------------------

@staff_member_required
def admin_area_regions(request):
    """List and manage area regions."""
    from .models import AreaRegion
    
    areas = AreaRegion.objects.all().order_by('region', 'area_name')
    
    # Group by region
    areas_by_region = {}
    for area in areas:
        region_display = area.get_region_display()
        if region_display not in areas_by_region:
            areas_by_region[region_display] = []
        areas_by_region[region_display].append(area)
    
    context = {
        'areas': areas,
        'areas_by_region': areas_by_region,
        'region_choices': AreaRegion.REGION_CHOICES,
        'total_areas': areas.count(),
    }
    return render(request, 'admin/admin_area_regions.html', context)


@staff_member_required
def admin_area_region_save(request):
    """Add or update an area region."""
    from .models import AreaRegion
    
    if request.method == 'POST':
        area_id = request.POST.get('id')
        area_name = request.POST.get('area_name', '').strip()
        region = request.POST.get('region', '').strip()
        pincode = request.POST.get('pincode', '').strip()
        is_active = request.POST.get('is_active') == 'on'
        
        if not area_name or not region:
            messages.error(request, 'Area name and region are required.')
            return redirect('admin_area_regions')
        
        try:
            if area_id:
                # Update existing
                area = AreaRegion.objects.get(id=area_id)
                area.area_name = area_name
                area.region = region
                area.pincode = pincode
                area.is_active = is_active
                area.save()
                messages.success(request, f'Area "{area_name}" updated successfully.')
            else:
                # Create new
                AreaRegion.objects.create(
                    area_name=area_name,
                    region=region,
                    pincode=pincode,
                    is_active=is_active
                )
                messages.success(request, f'Area "{area_name}" added successfully.')
        except Exception as e:
            messages.error(request, f'Error: {str(e)}')
    
    return redirect('admin_area_regions')


@staff_member_required
def admin_area_region_delete(request, pk):
    """Delete an area region."""
    from .models import AreaRegion
    
    try:
        area = AreaRegion.objects.get(id=pk)
        name = area.area_name
        area.delete()
        messages.success(request, f'Area "{name}" deleted successfully.')
    except AreaRegion.DoesNotExist:
        messages.error(request, 'Area not found.')
    
    return redirect('admin_area_regions')


# ------------------------------------------------------------------
# Site Announcements Management
# ------------------------------------------------------------------

@staff_member_required
def admin_announcements_list(request):
    """List all site announcements."""
    from .models import SiteAnnouncement
    
    announcements = SiteAnnouncement.objects.all().order_by('-created_at')
    
    # Count active announcements
    active_count = announcements.filter(is_active=True).count()
    
    return render(request, 'admin/admin_announcements_list.html', {
        'announcements': announcements,
        'active_count': active_count
    })


@staff_member_required
def admin_announcement_create(request):
    """Create a new site announcement."""
    from .models import SiteAnnouncement
    
    if request.method == 'POST':
        title = request.POST.get('title')
        message = request.POST.get('message')
        position = request.POST.get('position', 'bottom_right')
        style = request.POST.get('style', 'info')
        target_audience = request.POST.get('target_audience', 'all')
        is_active = request.POST.get('is_active') == 'on'
        is_dismissible = request.POST.get('is_dismissible') == 'on'
        show_once_per_session = request.POST.get('show_once_per_session') == 'on'
        button_text = request.POST.get('button_text') or None
        button_url = request.POST.get('button_url') or None
        start_date = request.POST.get('start_date') or None
        end_date = request.POST.get('end_date') or None
        
        announcement = SiteAnnouncement.objects.create(
            title=title,
            message=message,
            position=position,
            style=style,
            target_audience=target_audience,
            is_active=is_active,
            is_dismissible=is_dismissible,
            show_once_per_session=show_once_per_session,
            button_text=button_text,
            button_url=button_url,
            start_date=start_date,
            end_date=end_date,
            created_by=request.user
        )
        
        AdminActivityLog.objects.create(
            admin=request.user,
            action='CREATE',
            module='ANNOUNCEMENT',
            description=f"Created announcement: {title}",
            ip_address=request.META.get('REMOTE_ADDR')
        )
        
        messages.success(request, f'Announcement "{title}" created successfully!')
        return redirect('admin_announcements_list')
    
    return render(request, 'admin/admin_announcement_form.html', {
        'mode': 'create',
        'position_choices': SiteAnnouncement.POSITION_CHOICES,
        'style_choices': SiteAnnouncement.STYLE_CHOICES,
        'target_choices': SiteAnnouncement.TARGET_CHOICES,
    })


@staff_member_required
def admin_announcement_edit(request, pk):
    """Edit an existing site announcement."""
    from .models import SiteAnnouncement
    
    announcement = get_object_or_404(SiteAnnouncement, pk=pk)
    
    if request.method == 'POST':
        announcement.title = request.POST.get('title')
        announcement.message = request.POST.get('message')
        announcement.position = request.POST.get('position', 'bottom_right')
        announcement.style = request.POST.get('style', 'info')
        announcement.target_audience = request.POST.get('target_audience', 'all')
        announcement.is_active = request.POST.get('is_active') == 'on'
        announcement.is_dismissible = request.POST.get('is_dismissible') == 'on'
        announcement.show_once_per_session = request.POST.get('show_once_per_session') == 'on'
        announcement.button_text = request.POST.get('button_text') or None
        announcement.button_url = request.POST.get('button_url') or None
        announcement.start_date = request.POST.get('start_date') or None
        announcement.end_date = request.POST.get('end_date') or None
        announcement.save()
        
        AdminActivityLog.objects.create(
            admin=request.user,
            action='UPDATE',
            module='ANNOUNCEMENT',
            description=f"Updated announcement: {announcement.title}",
            ip_address=request.META.get('REMOTE_ADDR')
        )
        
        messages.success(request, f'Announcement "{announcement.title}" updated successfully!')
        return redirect('admin_announcements_list')
    
    return render(request, 'admin/admin_announcement_form.html', {
        'mode': 'edit',
        'announcement': announcement,
        'position_choices': SiteAnnouncement.POSITION_CHOICES,
        'style_choices': SiteAnnouncement.STYLE_CHOICES,
        'target_choices': SiteAnnouncement.TARGET_CHOICES,
    })


@staff_member_required
def admin_announcement_delete(request, pk):
    """Delete a site announcement."""
    from .models import SiteAnnouncement
    
    announcement = get_object_or_404(SiteAnnouncement, pk=pk)
    title = announcement.title
    announcement.delete()
    
    AdminActivityLog.objects.create(
        admin=request.user,
        action='DELETE',
        module='ANNOUNCEMENT',
        description=f"Deleted announcement: {title}",
        ip_address=request.META.get('REMOTE_ADDR')
    )
    
    messages.success(request, f'Announcement "{title}" deleted successfully!')
    return redirect('admin_announcements_list')


@staff_member_required
def admin_announcement_toggle(request, pk):
    """Toggle announcement active status."""
    from .models import SiteAnnouncement
    
    announcement = get_object_or_404(SiteAnnouncement, pk=pk)
    announcement.is_active = not announcement.is_active
    announcement.save()
    
    status = "activated" if announcement.is_active else "deactivated"
    messages.success(request, f'Announcement "{announcement.title}" {status}!')
    
    return redirect('admin_announcements_list')


# ==============================================================================
# ELECTRICIAN MANAGEMENT
# ==============================================================================

@staff_required
def admin_electrician_list(request):
    """List all electricians with search and filter options."""
    electricians = Electrician.objects.all()
    
    # Search functionality
    search = request.GET.get('search', '')
    if search:
        electricians = electricians.filter(
            models.Q(name__icontains=search) |
            models.Q(phone_number__icontains=search) |
            models.Q(email__icontains=search)
        )
    
    # Filter by status
    status_filter = request.GET.get('status', '')
    if status_filter == 'active':
        electricians = electricians.filter(is_active=True)
    elif status_filter == 'inactive':
        electricians = electricians.filter(is_active=False)
    elif status_filter == 'visible':
        electricians = electricians.filter(show_on_home_page=True, is_active=True)
    
    # Pagination
    paginator = Paginator(electricians, 20)
    page = request.GET.get('page', 1)
    electricians = paginator.get_page(page)
    
    return render(request, 'admin/admin_electrician_list.html', {
        'electricians': electricians,
        'search': search,
        'status_filter': status_filter,
    })


@staff_required
def admin_electrician_add(request):
    """Add a new electrician."""
    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        phone = request.POST.get('phone_number', '').strip()
        email = request.POST.get('email', '').strip()
        specializations = request.POST.get('specializations', '').strip()
        experience = request.POST.get('experience_years', 0)
        show_on_home = request.POST.get('show_on_home_page') == 'on'
        is_active = request.POST.get('is_active', 'on') == 'on'
        address = request.POST.get('address', '').strip()
        notes = request.POST.get('notes', '').strip()
        
        if not name or not phone or not email:
            messages.error(request, 'Name, Phone, and Email are required.')
            return redirect('admin_electrician_add')
        
        electrician = Electrician.objects.create(
            name=name,
            phone_number=phone,
            email=email,
            specializations=specializations,
            experience_years=int(experience) if experience else 0,
            show_on_home_page=show_on_home,
            is_active=is_active,
            address=address,
            notes=notes,
        )
        
        # Handle profile picture
        if 'profile_picture' in request.FILES:
            electrician.profile_picture = request.FILES['profile_picture']
            electrician.save()
        
        AdminActivityLog.objects.create(
            admin=request.user,
            action='CREATE',
            module='ELECTRICIAN',
            description=f"Added electrician: {name}",
            ip_address=request.META.get('REMOTE_ADDR')
        )
        
        messages.success(request, f'Electrician "{name}" added successfully!')
        return redirect('admin_electrician_list')
    
    return render(request, 'admin/admin_electrician_form.html', {
        'mode': 'add',
    })


@staff_required
def admin_electrician_edit(request, pk):
    """Edit an existing electrician."""
    electrician = get_object_or_404(Electrician, pk=pk)
    
    if request.method == 'POST':
        electrician.name = request.POST.get('name', '').strip()
        electrician.phone_number = request.POST.get('phone_number', '').strip()
        electrician.email = request.POST.get('email', '').strip()
        electrician.specializations = request.POST.get('specializations', '').strip()
        electrician.experience_years = int(request.POST.get('experience_years', 0) or 0)
        electrician.show_on_home_page = request.POST.get('show_on_home_page') == 'on'
        electrician.is_active = request.POST.get('is_active', 'on') == 'on'
        electrician.address = request.POST.get('address', '').strip()
        electrician.notes = request.POST.get('notes', '').strip()
        
        if 'profile_picture' in request.FILES:
            electrician.profile_picture = request.FILES['profile_picture']
        
        electrician.save()
        
        AdminActivityLog.objects.create(
            admin=request.user,
            action='UPDATE',
            module='ELECTRICIAN',
            description=f"Updated electrician: {electrician.name}",
            ip_address=request.META.get('REMOTE_ADDR')
        )
        
        messages.success(request, f'Electrician "{electrician.name}" updated successfully!')
        return redirect('admin_electrician_list')
    
    return render(request, 'admin/admin_electrician_form.html', {
        'mode': 'edit',
        'electrician': electrician,
    })


@staff_required
def admin_electrician_delete(request, pk):
    """Delete an electrician."""
    electrician = get_object_or_404(Electrician, pk=pk)
    name = electrician.name
    electrician.delete()
    
    AdminActivityLog.objects.create(
        admin=request.user,
        action='DELETE',
        module='ELECTRICIAN',
        description=f"Deleted electrician: {name}",
        ip_address=request.META.get('REMOTE_ADDR')
    )
    
    messages.success(request, f'Electrician "{name}" deleted successfully!')
    return redirect('admin_electrician_list')


@staff_required
def admin_electrician_toggle_visibility(request, pk):
    """Toggle electrician visibility on home page."""
    electrician = get_object_or_404(Electrician, pk=pk)
    electrician.show_on_home_page = not electrician.show_on_home_page
    electrician.save()
    
    status = "visible on home page" if electrician.show_on_home_page else "hidden from home page"
    messages.success(request, f'Electrician "{electrician.name}" is now {status}.')
    return redirect('admin_electrician_list')


# ==============================================================================
# WARRANTY MANAGEMENT
# ==============================================================================

@staff_required
def admin_warranty_list(request):
    """List all warranties with search and filter options."""
    warranties = Warranty.objects.all().select_related('customer', 'created_by')
    
    # Update expired warranties
    Warranty.update_expired_warranties()
    
    # Search functionality
    search = request.GET.get('search', '')
    if search:
        warranties = warranties.filter(
            models.Q(customer_name__icontains=search) |
            models.Q(customer_email__icontains=search) |
            models.Q(product_name__icontains=search) |
            models.Q(product_serial__icontains=search)
        )
    
    # Filter by status
    status_filter = request.GET.get('status', '')
    if status_filter:
        warranties = warranties.filter(status=status_filter)
    
    # Pagination
    paginator = Paginator(warranties, 20)
    page = request.GET.get('page', 1)
    warranties = paginator.get_page(page)
    
    return render(request, 'admin/admin_warranty_list.html', {
        'warranties': warranties,
        'search': search,
        'status_filter': status_filter,
        'status_choices': Warranty.STATUS_CHOICES,
    })


@staff_required
def admin_warranty_add(request):
    """Add a new warranty."""
    if request.method == 'POST':
        customer_name = request.POST.get('customer_name', '').strip()
        customer_phone = request.POST.get('customer_phone', '').strip()
        customer_email = request.POST.get('customer_email', '').strip()
        product_name = request.POST.get('product_name', '').strip()
        product_serial = request.POST.get('product_serial', '').strip()
        product_brand = request.POST.get('product_brand', '').strip()
        product_model = request.POST.get('product_model', '').strip()
        purchase_date = request.POST.get('purchase_date')
        warranty_duration = request.POST.get('warranty_duration', 12)
        warranty_unit = request.POST.get('warranty_unit', 'MONTHS')
        purchase_invoice = request.POST.get('purchase_invoice', '').strip()
        notes = request.POST.get('notes', '').strip()
        
        if not customer_name or not customer_phone or not customer_email or not product_name or not purchase_date:
            messages.error(request, 'Customer name, phone, email, product name, and purchase date are required.')
            return redirect('admin_warranty_add')
        
        warranty = Warranty.objects.create(
            customer_name=customer_name,
            customer_phone=customer_phone,
            customer_email=customer_email,
            product_name=product_name,
            product_serial=product_serial or None,
            product_brand=product_brand or None,
            product_model=product_model or None,
            purchase_date=purchase_date,
            warranty_duration=int(warranty_duration),
            warranty_unit=warranty_unit,
            purchase_invoice=purchase_invoice or None,
            notes=notes or None,
            created_by=request.user,
        )
        
        # Handle product image
        if 'product_image' in request.FILES:
            warranty.product_image = request.FILES['product_image']
            warranty.save()
        
        AdminActivityLog.objects.create(
            admin=request.user,
            action='CREATE',
            module='WARRANTY',
            description=f"Added warranty for {product_name} - {customer_name}",
            ip_address=request.META.get('REMOTE_ADDR')
        )
        
        messages.success(request, f'Warranty for "{product_name}" added successfully!')
        return redirect('admin_warranty_list')
    
    return render(request, 'admin/admin_warranty_form.html', {
        'mode': 'add',
        'unit_choices': Warranty.DURATION_UNIT_CHOICES,
    })


@staff_required
def admin_warranty_edit(request, pk):
    """Edit an existing warranty."""
    warranty = get_object_or_404(Warranty, pk=pk)
    
    if request.method == 'POST':
        warranty.customer_name = request.POST.get('customer_name', '').strip()
        warranty.customer_phone = request.POST.get('customer_phone', '').strip()
        warranty.customer_email = request.POST.get('customer_email', '').strip()
        warranty.product_name = request.POST.get('product_name', '').strip()
        warranty.product_serial = request.POST.get('product_serial', '').strip() or None
        warranty.product_brand = request.POST.get('product_brand', '').strip() or None
        warranty.product_model = request.POST.get('product_model', '').strip() or None
        warranty.purchase_date = request.POST.get('purchase_date')
        warranty.warranty_duration = int(request.POST.get('warranty_duration', 12))
        warranty.warranty_unit = request.POST.get('warranty_unit', 'MONTHS')
        warranty.purchase_invoice = request.POST.get('purchase_invoice', '').strip() or None
        warranty.notes = request.POST.get('notes', '').strip() or None
        
        if 'product_image' in request.FILES:
            warranty.product_image = request.FILES['product_image']
        
        warranty.save()
        
        AdminActivityLog.objects.create(
            admin=request.user,
            action='UPDATE',
            module='WARRANTY',
            description=f"Updated warranty for {warranty.product_name}",
            ip_address=request.META.get('REMOTE_ADDR')
        )
        
        messages.success(request, f'Warranty for "{warranty.product_name}" updated successfully!')
        return redirect('admin_warranty_list')
    
    return render(request, 'admin/admin_warranty_form.html', {
        'mode': 'edit',
        'warranty': warranty,
        'unit_choices': Warranty.DURATION_UNIT_CHOICES,
    })


@staff_required
def admin_warranty_void(request, pk):
    """Void a warranty."""
    warranty = get_object_or_404(Warranty, pk=pk)
    
    if request.method == 'POST':
        reason = request.POST.get('void_reason', 'Voided by admin').strip()
        warranty.void_warranty(reason, request.user)
        
        AdminActivityLog.objects.create(
            admin=request.user,
            action='UPDATE',
            module='WARRANTY',
            description=f"Voided warranty for {warranty.product_name}: {reason}",
            ip_address=request.META.get('REMOTE_ADDR')
        )
        
        messages.success(request, f'Warranty for "{warranty.product_name}" has been voided.')
        return redirect('admin_warranty_list')
    
    return render(request, 'admin/admin_warranty_void.html', {
        'warranty': warranty,
    })


@staff_required  
def admin_warranty_detail(request, pk):
    """View warranty details."""
    warranty = get_object_or_404(Warranty, pk=pk)
    return render(request, 'admin/admin_warranty_detail.html', {
        'warranty': warranty,
    })


# ==============================================================================
# ADMIN-BOOKED APPOINTMENTS
# ==============================================================================

@staff_required
def admin_appointment_book(request):
    """Book an appointment on behalf of a customer."""
    from .email_utils import send_professional_email
    
    if request.method == 'POST':
        customer_name = request.POST.get('customer_name', '').strip()
        phone = request.POST.get('phone', '').strip()
        email = request.POST.get('email', '').strip()
        service_id = request.POST.get('service')
        electrician_id = request.POST.get('electrician')
        date = request.POST.get('date')
        time = request.POST.get('time')
        visiting_charge = request.POST.get('visiting_charge', 200)
        notes = request.POST.get('notes', '').strip()
        
        # Address fields
        house_number = request.POST.get('house_number', '').strip()
        address_line1 = request.POST.get('address_line1', '').strip()
        address_line2 = request.POST.get('address_line2', '').strip()
        area = request.POST.get('area', 'Other')
        pincode = request.POST.get('pincode', '452001').strip()
        
        if not customer_name or not phone or not date or not time:
            messages.error(request, 'Customer name, phone, date, and time are required.')
            return redirect('admin_appointment_book')
        
        # Create appointment
        appointment = Appointment.objects.create(
            customer_name=customer_name,
            phone=phone,
            email=email or '',
            house_number=house_number,
            address_line1=address_line1,
            address_line2=address_line2,
            area=area,
            pincode=pincode,
            city='Indore',
            date=date,
            time=time,
            visiting_charge=float(visiting_charge) if visiting_charge else 200,
            problem_description=notes or 'Admin booked appointment',
            is_admin_booked=True,
            admin_notes=notes,
            status='Confirmed',
        )
        
        # Assign service
        if service_id:
            try:
                service = ServiceType.objects.get(pk=service_id)
                appointment.service = service
            except ServiceType.DoesNotExist:
                pass
        
        # Assign electrician
        if electrician_id:
            try:
                electrician = Electrician.objects.get(pk=electrician_id)
                appointment.assigned_electrician = electrician
            except Electrician.DoesNotExist:
                pass
        
        appointment.save()
        
        # Send email to customer if email provided
        if email:
            try:
                context = {
                    'customer_name': customer_name,
                    'appointment': appointment,
                    'electrician': appointment.assigned_electrician,
                    'business_profile_link': 'https://maps.google.com/?cid=your_business_id',
                }
                send_professional_email(
                    email_type='APPOINTMENT_STATUS',
                    recipient=email,
                    subject='Appointment Confirmed - Shiv Shakti Electricals',
                    template_name='emails/admin_booked_appointment.html',
                    context=context,
                    appointment=appointment,
                )
            except Exception as e:
                pass  # Email failed but appointment created
        
        # Send email to electrician if assigned
        if appointment.assigned_electrician and appointment.assigned_electrician.email:
            try:
                elec_context = {
                    'electrician': appointment.assigned_electrician,
                    'appointment': appointment,
                    'customer_name': customer_name,
                    'customer_phone': phone,
                    'address': f"{house_number}, {address_line1}, {area}, {pincode}",
                }
                send_professional_email(
                    email_type='APPOINTMENT_STATUS',
                    recipient=appointment.assigned_electrician.email,
                    subject=f'New Appointment Assignment - {customer_name}',
                    template_name='emails/electrician_assignment.html',
                    context=elec_context,
                    appointment=appointment,
                )
            except Exception as e:
                pass
        
        AdminActivityLog.objects.create(
            admin=request.user,
            action='CREATE',
            module='APPOINTMENT',
            description=f"Admin booked appointment for {customer_name}",
            ip_address=request.META.get('REMOTE_ADDR')
        )
        
        messages.success(request, f'Appointment for "{customer_name}" booked successfully!')
        return redirect('admin_appointment_list')
    
    services = ServiceType.get_active_services()
    electricians = Electrician.get_active_electricians()
    areas = Appointment.AREA_CHOICES
    
    return render(request, 'admin/admin_appointment_book.html', {
        'services': services,
        'electricians': electricians,
        'areas': areas,
    })


# ==============================================================================
# PRODUCT VISIBILITY CONTROLS
# ==============================================================================

@staff_required
def admin_product_toggle_visibility(request, pk):
    """Toggle product visibility on website."""
    product = get_object_or_404(Product, pk=pk)
    product.is_visible_on_website = not product.is_visible_on_website
    product.save()
    
    status = "visible on website" if product.is_visible_on_website else "hidden from website"
    messages.success(request, f'Product "{product.name}" is now {status}.')
    
    return redirect('admin_product_list')
